#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
关键词格式化模块
实现[关键词]黑体方括号格式
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

class KeywordFormatter:
    """关键词格式化器"""
    
    def format_keywords(self, doc, structure):
        """
        格式化中英文关键词
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        # 处理中文关键词
        cn_abstract_idx = structure.get('abstract_cn', -1)
        if cn_abstract_idx >= 0:
            self._format_cn_keywords(doc, cn_abstract_idx)
        
        # 处理英文关键词
        en_abstract_idx = structure.get('abstract_en', -1)
        if en_abstract_idx >= 0:
            self._format_en_keywords(doc, en_abstract_idx)
    
    def _format_cn_keywords(self, doc, start_idx):
        """格式化中文关键词"""
        # 在摘要后的10个段落内查找关键词
        for i in range(start_idx, min(start_idx + 15, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # 查找包含"关键词"的段落
            if self._contains_keywords_cn(text):
                # 提取关键词内容
                keywords = self._extract_keywords_cn(text)
                
                # 清空段落并重新格式化
                para.clear()
                
                # 添加[关键词]标签
                run_label = para.add_run('[关键词]')
                run_label.font.name = '黑体'
                run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run_label.font.size = Pt(12)  # 小四号
                run_label.font.bold = True
                
                # 添加空格
                para.add_run(' ')
                
                # 添加关键词内容
                run_content = para.add_run(keywords)
                run_content.font.name = '楷体'
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                run_content.font.size = Pt(12)  # 小四号
                
                # 设置段落格式
                para.paragraph_format.line_spacing = Pt(22)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                break
    
    def _format_en_keywords(self, doc, start_idx):
        """格式化英文关键词"""
        # 在摘要后的10个段落内查找关键词
        for i in range(start_idx, min(start_idx + 15, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # 查找包含"Keywords"的段落
            if self._contains_keywords_en(text):
                # 提取关键词内容
                keywords = self._extract_keywords_en(text)
                
                # 清空段落并重新格式化
                para.clear()
                
                # 添加[Keywords]标签
                run_label = para.add_run('[Keywords]')
                run_label.font.name = 'Arial Black'
                run_label.font.size = Pt(12)  # 小四号
                run_label.font.bold = True
                
                # 添加空格
                para.add_run(' ')
                
                # 添加关键词内容
                run_content = para.add_run(keywords)
                run_content.font.name = 'Times New Roman'
                run_content.font.size = Pt(12)  # 小四号
                
                # 设置段落格式
                para.paragraph_format.line_spacing = Pt(22)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                break
    
    def _contains_keywords_cn(self, text):
        """检查是否包含中文关键词标记"""
        keywords_patterns = [
            '关键词',
            '关键字',
            '【关键词】',
            '[关键词]',
            '关键词：',
            '关键词:'
        ]
        return any(pattern in text for pattern in keywords_patterns)
    
    def _contains_keywords_en(self, text):
        """检查是否包含英文关键词标记"""
        text_lower = text.lower()
        keywords_patterns = [
            'keywords',
            'key words',
            '[keywords]',
            '【keywords】',
            'keywords:',
            'keywords：'
        ]
        return any(pattern in text_lower for pattern in keywords_patterns)
    
    def _extract_keywords_cn(self, text):
        """提取中文关键词内容"""
        # 移除各种标记
        patterns_to_remove = [
            r'\[关键词\]',
            r'【关键词】',
            r'关键词[:：]',
            r'关键字[:：]',
            r'^\s*关键词\s*',
            r'^\s*关键字\s*'
        ]
        
        keywords = text
        for pattern in patterns_to_remove:
            keywords = re.sub(pattern, '', keywords, flags=re.IGNORECASE)
        
        # 清理空白字符
        keywords = keywords.strip()
        
        # 确保关键词之间用分号分隔
        if '，' in keywords and '；' not in keywords:
            keywords = keywords.replace('，', '；')
        elif ',' in keywords and ';' not in keywords and not any(ord(c) > 127 for c in keywords):
            keywords = keywords.replace(',', ';')
        
        return keywords
    
    def _extract_keywords_en(self, text):
        """提取英文关键词内容"""
        # 移除各种标记
        patterns_to_remove = [
            r'\[keywords\]',
            r'\[Keywords\]',
            r'【keywords】',
            r'【Keywords】',
            r'keywords[:：]',
            r'Keywords[:：]',
            r'key\s*words[:：]',
            r'^\s*keywords\s*',
            r'^\s*Keywords\s*'
        ]
        
        keywords = text
        for pattern in patterns_to_remove:
            keywords = re.sub(pattern, '', keywords, flags=re.IGNORECASE)
        
        # 清理空白字符
        keywords = keywords.strip()
        
        # 确保关键词之间用分号分隔
        if ',' in keywords and ';' not in keywords:
            keywords = keywords.replace(',', ';')
        
        return keywords
    
    def add_keywords_if_missing(self, doc, structure, keywords_dict=None):
        """
        如果缺少关键词，添加关键词
        
        Args:
            doc: Document对象
            structure: 文档结构字典
            keywords_dict: 包含'cn'和'en'键的字典，提供关键词内容
        """
        if keywords_dict is None:
            keywords_dict = {
                'cn': '深度学习；图像识别；卷积神经网络；特征提取；分类算法',
                'en': 'Deep Learning; Image Recognition; Convolutional Neural Network; Feature Extraction; Classification Algorithm'
            }
        
        # 检查并添加中文关键词
        cn_abstract_idx = structure.get('abstract_cn', -1)
        if cn_abstract_idx >= 0 and not self._has_keywords_cn(doc, cn_abstract_idx):
            self._add_keywords_cn(doc, cn_abstract_idx, keywords_dict.get('cn', ''))
        
        # 检查并添加英文关键词
        en_abstract_idx = structure.get('abstract_en', -1)
        if en_abstract_idx >= 0 and not self._has_keywords_en(doc, en_abstract_idx):
            self._add_keywords_en(doc, en_abstract_idx, keywords_dict.get('en', ''))
    
    def _has_keywords_cn(self, doc, start_idx):
        """检查是否已有中文关键词"""
        for i in range(start_idx, min(start_idx + 15, len(doc.paragraphs))):
            if self._contains_keywords_cn(doc.paragraphs[i].text):
                return True
        return False
    
    def _has_keywords_en(self, doc, start_idx):
        """检查是否已有英文关键词"""
        for i in range(start_idx, min(start_idx + 15, len(doc.paragraphs))):
            if self._contains_keywords_en(doc.paragraphs[i].text):
                return True
        return False
    
    def _add_keywords_cn(self, doc, abstract_idx, keywords):
        """在中文摘要后添加关键词"""
        # 找到摘要内容的最后一个段落
        insert_idx = abstract_idx + 1
        for i in range(abstract_idx + 1, min(abstract_idx + 10, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            # 如果遇到英文摘要或其他标题，停止
            if 'Abstract' in para.text or self._is_heading(para):
                break
            # 如果段落有内容，更新插入位置
            if para.text.strip():
                insert_idx = i + 1
        
        # 在适当位置插入关键词段落
        if insert_idx < len(doc.paragraphs):
            # 插入新段落
            new_para = doc.paragraphs[insert_idx - 1].insert_paragraph_after()
        else:
            new_para = doc.add_paragraph()
        
        # 格式化关键词
        run_label = new_para.add_run('[关键词]')
        run_label.font.name = '黑体'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_label.font.size = Pt(12)
        run_label.font.bold = True
        
        new_para.add_run(' ')
        
        run_content = new_para.add_run(keywords)
        run_content.font.name = '楷体'
        run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run_content.font.size = Pt(12)
    
    def _add_keywords_en(self, doc, abstract_idx, keywords):
        """在英文摘要后添加关键词"""
        # 找到摘要内容的最后一个段落
        insert_idx = abstract_idx + 1
        for i in range(abstract_idx + 1, min(abstract_idx + 10, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            # 如果遇到其他标题，停止
            if self._is_heading(para) and 'Abstract' not in para.text:
                break
            # 如果段落有内容，更新插入位置
            if para.text.strip():
                insert_idx = i + 1
        
        # 在适当位置插入关键词段落
        if insert_idx < len(doc.paragraphs):
            # 插入新段落
            new_para = doc.paragraphs[insert_idx - 1].insert_paragraph_after()
        else:
            new_para = doc.add_paragraph()
        
        # 格式化关键词
        run_label = new_para.add_run('[Keywords]')
        run_label.font.name = 'Arial Black'
        run_label.font.size = Pt(12)
        run_label.font.bold = True
        
        new_para.add_run(' ')
        
        run_content = new_para.add_run(keywords)
        run_content.font.name = 'Times New Roman'
        run_content.font.size = Pt(12)
    
    def _is_heading(self, paragraph):
        """判断是否为标题段落"""
        # 简单判断：字体较大或加粗
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font.bold or (first_run.font.size and first_run.font.size > Pt(12)):
                return True
        return False


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加中文摘要
    p1 = doc.add_paragraph("摘要")
    p1.runs[0].font.bold = True
    
    doc.add_paragraph("本文研究了基于深度学习的图像识别技术...")
    doc.add_paragraph("关键词：深度学习，图像识别，卷积神经网络")
    
    # 添加英文摘要
    p2 = doc.add_paragraph("Abstract")
    p2.runs[0].font.bold = True
    
    doc.add_paragraph("This paper studies image recognition technology based on deep learning...")
    doc.add_paragraph("Keywords: Deep Learning, Image Recognition, CNN")
    
    # 测试格式化
    formatter = KeywordFormatter()
    structure = {
        'abstract_cn': 0,
        'abstract_en': 3
    }
    
    formatter.format_keywords(doc, structure)
    
    # 保存测试文档
    doc.save('test_keywords.docx')
    print("测试文档已生成：test_keywords.docx")
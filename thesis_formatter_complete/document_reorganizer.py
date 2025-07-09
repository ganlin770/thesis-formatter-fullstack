#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档结构重组器模块
按照毕业论文装订顺序重新组织文档结构
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from copy import deepcopy

class DocumentReorganizer:
    """文档结构重组器"""
    
    def __init__(self):
        # 标准装订顺序
        self.standard_order = [
            'cover',           # 封面
            'commitment',      # 诚信承诺书
            'abstract_cn',     # 中文摘要
            'abstract_en',     # 英文摘要
            'toc',            # 目录
            'main_content',   # 正文（包括所有章节）
            'references',     # 参考文献
            'acknowledgment', # 致谢
            'appendix'        # 附录
        ]
        
        # 段落内容标识模式
        self.section_patterns = {
            'commitment': [r'诚信承诺', r'承诺书'],
            'abstract_cn': [r'^摘\s*要$', r'^摘要$'],
            'abstract_en': [r'^Abstract$', r'^ABSTRACT$'],
            'toc': [r'^目\s*录$', r'^目录$'],
            'references': [r'^参考文献$', r'^参\s*考\s*文\s*献$'],
            'acknowledgment': [r'^致\s*谢$', r'^致谢$', r'^谢\s*辞$'],
            'appendix': [r'^附录', r'^附\s*录', r'^Appendix', r'^APPENDIX']
        }
    
    def reorganize_document(self, doc, structure):
        """
        重组文档结构
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        
        Returns:
            重组后的Document对象
        """
        # 分析现有结构
        sections = self._analyze_document_structure(doc, structure)
        
        # 创建新文档
        new_doc = Document()
        
        # 复制文档样式
        self._copy_styles(doc, new_doc)
        
        # 按标准顺序重组内容
        for section_name in self.standard_order:
            if section_name in sections and sections[section_name]:
                self._copy_section_to_new_doc(
                    doc, 
                    new_doc, 
                    sections[section_name]['start'],
                    sections[section_name]['end'],
                    section_name
                )
        
        # 处理未识别的内容
        self._handle_unidentified_content(doc, new_doc, sections)
        
        return new_doc
    
    def _analyze_document_structure(self, doc, structure):
        """分析文档结构，返回各部分的位置"""
        sections = {}
        
        # 使用structure中已有的信息
        if 'cover' in structure and structure['cover'] >= 0:
            sections['cover'] = {
                'start': 0,
                'end': structure.get('cover_end', 0)
            }
        
        # 扫描文档识别各部分
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 检查每种类型的部分
            for section_type, patterns in self.section_patterns.items():
                for pattern in patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        # 找到部分的开始
                        if section_type not in sections:
                            sections[section_type] = {'start': i}
                        
                        # 更新前一个部分的结束位置
                        for sec_name, sec_info in sections.items():
                            if 'end' not in sec_info and sec_info['start'] < i:
                                sec_info['end'] = i - 1
                        break
        
        # 处理正文部分
        main_start = structure.get('main_start', -1)
        if main_start >= 0:
            # 找到正文结束位置（参考文献前）
            main_end = len(doc.paragraphs) - 1
            for section_type in ['references', 'acknowledgment', 'appendix']:
                if section_type in sections:
                    main_end = min(main_end, sections[section_type]['start'] - 1)
            
            sections['main_content'] = {
                'start': main_start,
                'end': main_end
            }
        
        # 设置最后一个部分的结束位置
        for sec_info in sections.values():
            if 'end' not in sec_info:
                sec_info['end'] = len(doc.paragraphs) - 1
        
        return sections
    
    def _copy_styles(self, source_doc, target_doc):
        """复制文档样式"""
        # 复制段落样式
        try:
            for style in source_doc.styles:
                if style.type == 1:  # 段落样式
                    if style.name not in [s.name for s in target_doc.styles]:
                        target_doc.styles.add_style(style.name, style.type)
        except:
            # 忽略样式复制错误
            pass
    
    def _copy_section_to_new_doc(self, source_doc, target_doc, start_idx, end_idx, section_name):
        """复制指定部分到新文档"""
        # 特殊处理某些部分
        if section_name == 'cover':
            # 封面不需要前置分页
            pass
        elif section_name in ['commitment', 'abstract_cn', 'toc', 'main_content', 'references', 'acknowledgment', 'appendix']:
            # 这些部分前需要分页
            if len(target_doc.paragraphs) > 0:
                target_doc.add_page_break()
        
        # 复制段落
        for i in range(start_idx, min(end_idx + 1, len(source_doc.paragraphs))):
            source_para = source_doc.paragraphs[i]
            
            # 创建新段落
            new_para = target_doc.add_paragraph()
            
            # 复制段落格式
            self._copy_paragraph_format(source_para, new_para)
            
            # 复制段落内容
            for run in source_para.runs:
                new_run = new_para.add_run(run.text)
                self._copy_run_format(run, new_run)
            
            # 处理分页符
            if self._has_page_break(source_para):
                target_doc.add_page_break()
    
    def _copy_paragraph_format(self, source_para, target_para):
        """复制段落格式"""
        fmt = source_para.paragraph_format
        target_fmt = target_para.paragraph_format
        
        # 复制对齐方式
        if source_para.alignment is not None:
            target_para.alignment = source_para.alignment
        
        # 复制缩进
        if fmt.first_line_indent is not None:
            target_fmt.first_line_indent = fmt.first_line_indent
        if fmt.left_indent is not None:
            target_fmt.left_indent = fmt.left_indent
        if fmt.right_indent is not None:
            target_fmt.right_indent = fmt.right_indent
        
        # 复制间距
        if fmt.space_before is not None:
            target_fmt.space_before = fmt.space_before
        if fmt.space_after is not None:
            target_fmt.space_after = fmt.space_after
        if fmt.line_spacing is not None:
            target_fmt.line_spacing = fmt.line_spacing
    
    def _copy_run_format(self, source_run, target_run):
        """复制文本格式"""
        # 复制字体
        if source_run.font.name:
            target_run.font.name = source_run.font.name
            target_run._element.rPr.rFonts.set(qn('w:eastAsia'), source_run.font.name)
        
        # 复制字号
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        
        # 复制样式
        if source_run.font.bold is not None:
            target_run.font.bold = source_run.font.bold
        if source_run.font.italic is not None:
            target_run.font.italic = source_run.font.italic
        if source_run.font.underline is not None:
            target_run.font.underline = source_run.font.underline
    
    def _has_page_break(self, paragraph):
        """检查段落是否包含分页符"""
        for run in paragraph.runs:
            if '\f' in run.text or '\x0c' in run.text:
                return True
            # 检查XML中的分页符
            for child in run._element:
                if child.tag.endswith('br'):
                    type_attr = child.get(qn('w:type'))
                    if type_attr and 'page' in type_attr:
                        return True
        return False
    
    def _handle_unidentified_content(self, source_doc, target_doc, sections):
        """处理未识别的内容"""
        # 收集所有已处理的段落索引
        processed_indices = set()
        for section_info in sections.values():
            for i in range(section_info['start'], section_info['end'] + 1):
                processed_indices.add(i)
        
        # 收集未处理的段落
        unprocessed = []
        for i, para in enumerate(source_doc.paragraphs):
            if i not in processed_indices and para.text.strip():
                unprocessed.append(i)
        
        # 如果有未处理的内容，添加到文档末尾
        if unprocessed:
            target_doc.add_page_break()
            p_warning = target_doc.add_paragraph()
            p_warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_warning = p_warning.add_run("【以下内容未能自动归类】")
            run_warning.font.name = '宋体'
            run_warning._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_warning.font.size = Pt(12)
            run_warning.font.bold = True
            
            for idx in unprocessed:
                source_para = source_doc.paragraphs[idx]
                new_para = target_doc.add_paragraph()
                self._copy_paragraph_format(source_para, new_para)
                
                for run in source_para.runs:
                    new_run = new_para.add_run(run.text)
                    self._copy_run_format(run, new_run)
    
    def validate_document_order(self, doc):
        """
        验证文档是否符合标准装订顺序
        
        Returns:
            (bool, str): (是否符合, 说明信息)
        """
        structure = self._analyze_document_structure(doc, {})
        
        issues = []
        
        # 检查必需部分
        required_sections = ['abstract_cn', 'toc', 'main_content', 'references']
        for section in required_sections:
            if section not in structure:
                issues.append(f"缺少{self._get_section_chinese_name(section)}")
        
        # 检查顺序
        found_sections = []
        for section in self.standard_order:
            if section in structure:
                found_sections.append(section)
        
        # 验证顺序是否正确
        last_pos = -1
        for section in found_sections:
            current_pos = structure[section]['start']
            if current_pos <= last_pos:
                issues.append(f"{self._get_section_chinese_name(section)}位置不正确")
            last_pos = current_pos
        
        if issues:
            return False, "装订顺序问题：" + "；".join(issues)
        else:
            return True, "文档装订顺序正确"
    
    def _get_section_chinese_name(self, section_key):
        """获取部分的中文名称"""
        name_map = {
            'cover': '封面',
            'commitment': '诚信承诺书',
            'abstract_cn': '中文摘要',
            'abstract_en': '英文摘要',
            'toc': '目录',
            'main_content': '正文',
            'references': '参考文献',
            'acknowledgment': '致谢',
            'appendix': '附录'
        }
        return name_map.get(section_key, section_key)


if __name__ == "__main__":
    # 测试代码
    from docx import Document
    
    # 创建测试文档（顺序混乱）
    doc = Document()
    
    # 先添加正文
    p1 = doc.add_paragraph("第一章 绪论")
    p1.runs[0].font.bold = True
    doc.add_paragraph("正文内容...")
    
    # 再添加摘要
    doc.add_page_break()
    p2 = doc.add_paragraph("摘要")
    p2.runs[0].font.bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("这是中文摘要...")
    doc.add_paragraph("关键词：测试，重组")
    
    # 添加参考文献
    doc.add_page_break()
    p3 = doc.add_paragraph("参考文献")
    p3.runs[0].font.bold = True
    doc.add_paragraph("[1] 测试文献")
    
    # 添加目录
    doc.add_page_break()
    p4 = doc.add_paragraph("目录")
    p4.runs[0].font.bold = True
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("第一章 绪论...1")
    
    # 保存原始文档
    doc.save('test_original_order.docx')
    
    # 测试重组
    reorganizer = DocumentReorganizer()
    structure = {
        'abstract_cn': 3,
        'toc': 8,
        'main_start': 0,
        'references': 6
    }
    
    # 验证原始顺序
    is_valid, message = reorganizer.validate_document_order(doc)
    print(f"原始文档验证: {message}")
    
    # 重组文档
    new_doc = reorganizer.reorganize_document(doc, structure)
    
    # 保存重组后的文档
    new_doc.save('test_reorganized.docx')
    print("重组后的文档已保存: test_reorganized.docx")
    
    # 验证重组后的顺序
    is_valid, message = reorganizer.validate_document_order(new_doc)
    print(f"重组后文档验证: {message}")
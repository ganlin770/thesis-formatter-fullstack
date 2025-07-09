#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
致谢格式化模块
实现标准致谢格式处理
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

class AcknowledgmentFormatter:
    """致谢格式化器"""
    
    def __init__(self):
        self.acknowledgment_patterns = [
            '致谢', '致謝', '鸣谢', '謝辞',
            'Acknowledgment', 'Acknowledgement', 'Acknowledgments'
        ]
    
    def format_acknowledgment(self, doc, structure):
        """
        格式化致谢部分
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        # 查找致谢部分
        ack_idx = self._find_acknowledgment(doc, structure)
        
        if ack_idx >= 0:
            # 格式化致谢
            self._format_acknowledgment_section(doc, ack_idx)
        else:
            # 如果需要，可以添加致谢模板
            pass
    
    def _find_acknowledgment(self, doc, structure):
        """查找致谢部分的位置"""
        # 通常在参考文献后
        ref_idx = structure.get('references', -1)
        start_idx = ref_idx + 1 if ref_idx >= 0 else 0
        
        # 搜索致谢标题
        for i in range(start_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if self._is_acknowledgment_title(para):
                return i
        
        return -1
    
    def _is_acknowledgment_title(self, paragraph):
        """判断是否为致谢标题"""
        text = paragraph.text.strip()
        
        # 检查是否匹配致谢标题
        for pattern in self.acknowledgment_patterns:
            if pattern in text and len(text) < 20:
                # 检查格式特征
                if paragraph.runs and paragraph.runs[0].font.bold:
                    return True
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return True
        
        return False
    
    def _format_acknowledgment_section(self, doc, ack_idx):
        """格式化致谢部分"""
        # 格式化标题
        title_para = doc.paragraphs[ack_idx]
        self._format_acknowledgment_title(title_para)
        
        # 格式化内容段落
        self._format_acknowledgment_content(doc, ack_idx + 1)
    
    def _format_acknowledgment_title(self, paragraph):
        """格式化致谢标题"""
        # 保存原文本
        text = paragraph.text.strip()
        
        # 清空并重新格式化
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(18)  # 二号
        run.font.bold = True
        
        # 设置段落间距
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.line_spacing = Pt(22)
    
    def _format_acknowledgment_content(self, doc, start_idx):
        """格式化致谢内容"""
        # 查找致谢内容的结束位置
        end_idx = self._find_acknowledgment_end(doc, start_idx)
        
        # 格式化每个段落
        for i in range(start_idx, end_idx):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                
                # 跳过空段落
                if not para.text.strip():
                    continue
                
                # 格式化段落
                self._format_content_paragraph(para)
    
    def _find_acknowledgment_end(self, doc, start_idx):
        """查找致谢内容的结束位置"""
        for i in range(start_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            
            # 检查是否为新的章节标题
            if self._is_section_title(para):
                return i
            
            # 检查是否有分页符
            for run in para.runs:
                if run._element.xml.find('w:br') != -1:
                    return i
        
        # 默认到文档结尾
        return len(doc.paragraphs)
    
    def _is_section_title(self, paragraph):
        """判断是否为章节标题"""
        text = paragraph.text.strip()
        
        # 检查常见的章节标题模式
        title_patterns = [
            r'^第[一二三四五六七八九十\d]+[章节]',
            r'^附录',
            r'^Appendix',
            r'^参考文献',
            r'^References'
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, text):
                return True
        
        # 检查格式特征
        if paragraph.runs and paragraph.runs[0].font.bold and len(text) < 30:
            return True
        
        return False
    
    def _format_content_paragraph(self, paragraph):
        """格式化致谢内容段落"""
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        # 格式化文本
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)  # 小四号
            run.font.bold = False  # 正文不加粗
    
    def add_acknowledgment_template(self, doc, position=None):
        """
        添加致谢模板
        
        Args:
            doc: Document对象
            position: 插入位置，None表示在文档末尾
        """
        # 添加致谢标题
        if position is not None and position < len(doc.paragraphs):
            title_para = doc.paragraphs[position].insert_paragraph_before()
        else:
            title_para = doc.add_paragraph()
        
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title_para.add_run("致谢")
        run_title.font.name = '宋体'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_title.font.size = Pt(18)
        run_title.font.bold = True
        
        # 添加致谢内容模板
        template_text = """时光荏苒，转眼间在江西财经大学现代经济管理学院的学习生涯即将结束。在这里，我要向所有给予我帮助和支持的人表达最真挚的谢意。

首先，我要特别感谢我的指导老师XXX教授。在论文的选题、研究和撰写过程中，X老师给予了我悉心的指导和无私的帮助。他/她渊博的学识、严谨的治学态度和高尚的人格魅力，深深地影响和激励着我。

其次，感谢学院的各位老师们。在四年的学习过程中，老师们的谆谆教诲让我在专业知识和综合素质方面都得到了很大的提升。

同时，我要感谢我的同学和朋友们。感谢你们在学习和生活中给予我的帮助和陪伴，让我的大学生活充满了欢声笑语。

最后，我要特别感谢我的家人。感谢你们一直以来的理解、支持和鼓励，是你们的爱让我能够安心完成学业。

在即将踏上新的人生征程之际，我将铭记母校的教诲，继续努力，不断进取，以优异的成绩回报母校、老师和所有关心我的人。"""
        
        # 分段添加内容
        for para_text in template_text.strip().split('\n\n'):
            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content_para.paragraph_format.first_line_indent = Pt(24)
            content_para.paragraph_format.line_spacing = Pt(22)
            
            run = content_para.add_run(para_text.strip())
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
    
    def check_acknowledgment_content(self, doc, ack_idx):
        """
        检查致谢内容是否符合规范
        
        Returns:
            (bool, str): (是否符合规范, 提示信息)
        """
        if ack_idx < 0:
            return False, "未找到致谢部分"
        
        # 查找内容结束位置
        end_idx = self._find_acknowledgment_end(doc, ack_idx + 1)
        
        # 计算字数
        total_chars = 0
        for i in range(ack_idx + 1, end_idx):
            if i < len(doc.paragraphs):
                total_chars += len(doc.paragraphs[i].text.strip())
        
        # 检查字数要求（通常300-800字）
        if total_chars < 200:
            return False, f"致谢内容过短（当前{total_chars}字，建议不少于300字）"
        elif total_chars > 1500:
            return False, f"致谢内容过长（当前{total_chars}字，建议不超过800字）"
        
        # 检查是否包含必要元素
        content = ''
        for i in range(ack_idx + 1, end_idx):
            if i < len(doc.paragraphs):
                content += doc.paragraphs[i].text
        
        required_elements = ['导师', '老师', '感谢']
        missing_elements = []
        
        for element in required_elements:
            if element not in content:
                missing_elements.append(element)
        
        if missing_elements:
            return False, f"致谢内容可能缺少必要元素：{', '.join(missing_elements)}"
        
        return True, "致谢格式符合规范"


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加一些前置内容
    doc.add_paragraph("第五章 结论")
    doc.add_paragraph("本研究的主要结论如下...")
    
    # 添加参考文献
    ref_para = doc.add_paragraph("参考文献")
    ref_para.runs[0].font.bold = True
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("[1] 张三. 论文题目[J]. 期刊名称, 2024.")
    
    # 添加致谢
    ack_para = doc.add_paragraph("致谢")
    ack_para.runs[0].font.bold = True
    
    doc.add_paragraph("在论文完成之际，我要向所有帮助过我的人表示感谢。")
    doc.add_paragraph("首先，感谢我的导师李教授，他的悉心指导让我受益匪浅。")
    doc.add_paragraph("其次，感谢实验室的同学们，感谢大家的帮助和支持。")
    doc.add_paragraph("最后，感谢我的家人，是你们的理解和支持让我能够顺利完成学业。")
    
    # 测试格式化
    formatter = AcknowledgmentFormatter()
    structure = {'references': 2}
    
    formatter.format_acknowledgment(doc, structure)
    
    # 检查致谢内容
    ack_idx = formatter._find_acknowledgment(doc, structure)
    is_valid, message = formatter.check_acknowledgment_content(doc, ack_idx)
    print(f"致谢检查结果：{message}")
    
    # 保存测试文档
    doc.save('test_acknowledgment.docx')
    print("测试文档已生成：test_acknowledgment.docx")
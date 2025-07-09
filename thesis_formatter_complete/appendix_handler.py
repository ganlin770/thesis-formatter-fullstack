#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
附录处理器模块
识别和格式化论文附录部分
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

class AppendixHandler:
    """附录处理器"""
    
    def __init__(self):
        # 附录标题模式
        self.appendix_patterns = [
            r'^附录[A-Z]?[:：]?\s*',
            r'^附\s*录[A-Z]?[:：]?\s*',
            r'^Appendix\s*[A-Z]?[:：]?\s*',
            r'^APPENDIX\s*[A-Z]?[:：]?\s*'
        ]
        
        # 附录编号格式
        self.appendix_labels = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
    
    def process_appendix(self, doc, structure):
        """
        处理文档中的附录
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        # 查找所有附录
        appendix_sections = self._find_appendix_sections(doc, structure)
        
        if not appendix_sections:
            return
        
        # 格式化每个附录
        for idx, (start_idx, end_idx, title) in enumerate(appendix_sections):
            self._format_appendix_section(doc, start_idx, end_idx, idx)
    
    def _find_appendix_sections(self, doc, structure):
        """查找文档中的附录部分"""
        appendix_sections = []
        current_appendix_start = -1
        
        # 从参考文献后开始查找
        ref_idx = structure.get('references', len(doc.paragraphs) - 10)
        
        for i in range(ref_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # 检查是否为附录标题
            if self._is_appendix_title(text):
                # 如果有前一个附录，记录其结束位置
                if current_appendix_start >= 0:
                    appendix_sections.append((
                        current_appendix_start, 
                        i - 1, 
                        doc.paragraphs[current_appendix_start].text.strip()
                    ))
                
                current_appendix_start = i
        
        # 记录最后一个附录
        if current_appendix_start >= 0:
            appendix_sections.append((
                current_appendix_start,
                len(doc.paragraphs) - 1,
                doc.paragraphs[current_appendix_start].text.strip()
            ))
        
        return appendix_sections
    
    def _is_appendix_title(self, text):
        """检查是否为附录标题"""
        if not text or len(text) > 50:
            return False
        
        for pattern in self.appendix_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _format_appendix_section(self, doc, start_idx, end_idx, appendix_num):
        """格式化单个附录部分"""
        # 格式化附录标题
        title_para = doc.paragraphs[start_idx]
        self._format_appendix_title(title_para, appendix_num)
        
        # 格式化附录内容
        for i in range(start_idx + 1, min(end_idx + 1, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 检查是否为子标题
            if self._is_appendix_subtitle(para):
                self._format_appendix_subtitle(para)
            else:
                # 普通段落格式化
                self._format_appendix_paragraph(para)
    
    def _format_appendix_title(self, paragraph, appendix_num):
        """格式化附录主标题"""
        # 获取附录编号
        if appendix_num < len(self.appendix_labels):
            label = self.appendix_labels[appendix_num]
        else:
            label = str(appendix_num + 1)
        
        # 提取原始标题内容
        original_text = paragraph.text.strip()
        
        # 移除已有的附录标记
        for pattern in self.appendix_patterns:
            original_text = re.sub(pattern, '', original_text, flags=re.IGNORECASE)
        
        # 清空段落
        paragraph.clear()
        
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(18)
        
        # 添加格式化的标题
        if original_text:
            run = paragraph.add_run(f"附录{label}  {original_text}")
        else:
            run = paragraph.add_run(f"附录{label}")
        
        # 设置字体
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 三号
        run.font.bold = True
    
    def _is_appendix_subtitle(self, paragraph):
        """检查是否为附录子标题"""
        text = paragraph.text.strip()
        
        # 子标题特征
        if len(text) < 50 and (
            paragraph.runs and paragraph.runs[0].font.bold or
            re.match(r'^[A-Z]\.\d+', text) or  # A.1 格式
            re.match(r'^[一二三四五六七八九十]+、', text)  # 中文编号
        ):
            return True
        
        return False
    
    def _format_appendix_subtitle(self, paragraph):
        """格式化附录子标题"""
        # 保持原有格式，只调整基本属性
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        
        # 确保字体
        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            if not run.font.size:
                run.font.size = Pt(14)  # 四号
    
    def _format_appendix_paragraph(self, paragraph):
        """格式化附录普通段落"""
        # 设置段落格式
        paragraph.paragraph_format.first_line_indent = Pt(28)  # 两字符缩进
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(0)
        
        # 设置字体
        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            if not run.font.size:
                run.font.size = Pt(12)  # 小四号
    
    def add_appendix(self, doc, title, content):
        """
        添加新的附录
        
        Args:
            doc: Document对象
            title: 附录标题
            content: 附录内容（可以是字符串或段落列表）
        """
        # 添加分页符
        doc.add_page_break()
        
        # 添加附录标题
        p_title = doc.add_paragraph()
        
        # 获取下一个附录编号
        existing_appendices = self._count_existing_appendices(doc)
        appendix_num = existing_appendices
        
        if appendix_num < len(self.appendix_labels):
            label = self.appendix_labels[appendix_num]
        else:
            label = str(appendix_num + 1)
        
        # 设置标题格式
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(f"附录{label}  {title}")
        run_title.font.name = '黑体'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        
        # 添加内容
        if isinstance(content, str):
            # 字符串内容，按段落分割
            for para_text in content.split('\n'):
                if para_text.strip():
                    p_content = doc.add_paragraph()
                    p_content.paragraph_format.first_line_indent = Pt(28)
                    run_content = p_content.add_run(para_text.strip())
                    run_content.font.name = '宋体'
                    run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run_content.font.size = Pt(12)
        elif isinstance(content, list):
            # 段落列表
            for para_content in content:
                p_content = doc.add_paragraph()
                p_content.paragraph_format.first_line_indent = Pt(28)
                run_content = p_content.add_run(str(para_content))
                run_content.font.name = '宋体'
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_content.font.size = Pt(12)
    
    def _count_existing_appendices(self, doc):
        """统计现有附录数量"""
        count = 0
        for para in doc.paragraphs:
            if self._is_appendix_title(para.text.strip()):
                count += 1
        return count


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加一些前置内容
    doc.add_paragraph("第五章 结论")
    doc.add_paragraph("本研究的主要结论...")
    doc.add_page_break()
    
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 研究方法[M]. 北京: 科学出版社, 2023.")
    doc.add_page_break()
    
    # 添加附录
    handler = AppendixHandler()
    
    # 手动添加的附录（未格式化）
    p1 = doc.add_paragraph("附录A 调查问卷")
    p1.runs[0].font.bold = True
    
    doc.add_paragraph("1. 您的年龄段：")
    doc.add_paragraph("   □ 18-25岁  □ 26-35岁  □ 36-45岁  □ 45岁以上")
    
    doc.add_paragraph("2. 您的教育程度：")
    doc.add_paragraph("   □ 高中  □ 本科  □ 硕士  □ 博士")
    
    # 使用handler添加新附录
    handler.add_appendix(
        doc,
        "实验数据",
        """实验组1数据：
温度: 25°C, 压力: 101.3kPa
结果: 成功

实验组2数据：
温度: 30°C, 压力: 101.3kPa
结果: 成功"""
    )
    
    # 格式化所有附录
    structure = {'references': 3}
    handler.process_appendix(doc, structure)
    
    # 保存测试文档
    doc.save('test_appendix.docx')
    print("测试文档已生成：test_appendix.docx")
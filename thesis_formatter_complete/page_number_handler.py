#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
页码处理器模块
实现罗马数字（前置部分）和阿拉伯数字（正文部分）的页码系统
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START as WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class PageNumberHandler:
    """页码处理器 - 处理罗马数字和阿拉伯数字页码"""
    
    def __init__(self, document):
        self.document = document
        
    def setup_page_numbers(self, structure):
        """
        设置整个文档的页码系统
        
        Args:
            structure: 文档结构字典，包含各部分的位置信息
        """
        # 确保至少有两个节
        self._ensure_sections(structure)
        
        # 第一节：封面到目录（罗马数字）
        self._setup_roman_section(0)
        
        # 第二节：正文开始（阿拉伯数字）
        if len(self.document.sections) > 1:
            self._setup_arabic_section(1)
    
    def _ensure_sections(self, structure):
        """确保文档有正确的分节"""
        # 如果只有一个节，需要在正文开始处添加分节符
        if len(self.document.sections) == 1:
            main_start = structure.get('main_start', -1)
            if main_start > 0 and main_start < len(self.document.paragraphs):
                # 在正文开始前插入分节符
                para = self.document.paragraphs[main_start - 1]
                # 使用add_section方法代替
                self.document.add_section()
    
    def _setup_roman_section(self, section_index):
        """设置罗马数字页码节"""
        if section_index >= len(self.document.sections):
            return
            
        section = self.document.sections[section_index]
        
        # 设置页脚
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()
        
        # 添加罗马数字页码
        self._add_page_number_field(footer_para, 'roman')
        
        # 设置该节不链接到前一节
        section.footer.is_linked_to_previous = False
        
        # 从第一页开始编号
        # section.start_type = WD_SECTION.NEW_PAGE
        # section.page_number_start = 1  # 罗马数字从I开始
    
    def _setup_arabic_section(self, section_index):
        """设置阿拉伯数字页码节"""
        if section_index >= len(self.document.sections):
            return
            
        section = self.document.sections[section_index]
        
        # 设置页脚
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()
        
        # 添加阿拉伯数字页码
        self._add_page_number_field(footer_para, 'arabic')
        
        # 设置该节不链接到前一节
        section.footer.is_linked_to_previous = False
        
        # 重新开始编号
        # section.start_type = WD_SECTION.NEW_PAGE
        # 设置起始页码为1
        self._set_start_page_number(section, 1)
    
    def _add_page_number_field(self, paragraph, number_format='arabic'):
        """
        添加页码域到段落
        
        Args:
            paragraph: 要添加页码的段落
            number_format: 页码格式 ('roman' 或 'arabic')
        """
        run = paragraph.add_run()
        
        # 创建域代码元素
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        
        if number_format == 'roman':
            # 小写罗马数字
            instrText.text = ' PAGE \\* roman '
        else:
            # 阿拉伯数字
            instrText.text = ' PAGE \\* arabic '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # 页码占位符
        t = OxmlElement('w:t')
        t.text = '1'
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # 将元素添加到run
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(t)
        run._r.append(fldChar3)
        
        # 设置字体
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)  # 五号字
    
    def _set_start_page_number(self, section, start_number):
        """设置节的起始页码"""
        sectPr = section._sectPr
        
        # 查找或创建pgNumType元素
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        
        # 设置起始页码
        pgNumType.set(qn('w:start'), str(start_number))
    
    def remove_existing_page_numbers(self):
        """移除现有的页码（如果有）"""
        for section in self.document.sections:
            # 清空页脚
            footer = section.footer
            for para in footer.paragraphs:
                para.clear()
            
            # 清空页眉（如果有页码）
            header = section.header  
            for para in header.paragraphs:
                # 检查是否包含页码域
                if 'PAGE' in para.text:
                    para.clear()
    
    def add_section_break_before_main(self, main_start_index):
        """在正文开始前添加分节符"""
        if main_start_index > 0 and main_start_index < len(self.document.paragraphs):
            # 获取正文前的最后一个段落
            para_before_main = self.document.paragraphs[main_start_index - 1]
            
            # 检查是否已有分节符
            if not self._has_section_break(para_before_main):
                # 添加分节符（新页）
                self.document.add_section()
    
    def _has_section_break(self, paragraph):
        """检查段落是否包含分节符"""
        for run in paragraph.runs:
            for child in run._element:
                if child.tag.endswith('br'):
                    type_attr = child.get(qn('w:type'))
                    if type_attr and 'page' in type_attr:
                        return True
        return False
    
    def set_page_margins(self):
        """设置页边距"""
        for section in self.document.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
            
            # A4纸张大小
            section.page_height = Cm(29.7)
            section.page_width = Cm(21)


# 辅助类：页码样式管理
class PageNumberStyle:
    """页码样式配置"""
    
    ROMAN_LOWER = 'roman'        # i, ii, iii
    ROMAN_UPPER = 'ROMAN'        # I, II, III
    ARABIC = 'arabic'            # 1, 2, 3
    LETTER_LOWER = 'alphabetic'  # a, b, c
    LETTER_UPPER = 'ALPHABETIC'  # A, B, C
    
    @staticmethod
    def get_format_code(style):
        """获取Word域代码格式"""
        format_map = {
            'roman': '\\* roman',
            'ROMAN': '\\* ROMAN', 
            'arabic': '\\* arabic',
            'alphabetic': '\\* alphabetic',
            'ALPHABETIC': '\\* ALPHABETIC'
        }
        return format_map.get(style, '\\* arabic')


if __name__ == "__main__":
    # 测试代码
    from docx.shared import Pt, Cm
    
    doc = Document()
    
    # 添加一些测试内容
    # 封面
    doc.add_paragraph("封面内容")
    doc.add_page_break()
    
    # 摘要
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是中文摘要内容...")
    doc.add_page_break()
    
    # 目录
    doc.add_paragraph("目录")
    doc.add_paragraph("1. 第一章")
    doc.add_paragraph("2. 第二章")
    
    # 添加分节符
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_SECTION.NEW_PAGE)
    
    # 正文
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("正文内容开始...")
    
    # 设置页码
    handler = PageNumberHandler(doc)
    structure = {
        'abstract_cn': 1,
        'toc': 3,
        'main_start': 6
    }
    
    handler.setup_page_numbers(structure)
    
    # 保存测试文档
    doc.save('test_page_numbers.docx')
    print("测试文档已生成：test_page_numbers.docx")
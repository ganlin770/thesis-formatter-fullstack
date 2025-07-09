#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
目录生成器模块
实现自动生成和更新目录，包含页码对齐
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT
try:
    from docx.enum.text import WD_TAB_LEADER
except ImportError:
    # For compatibility
    class WD_TAB_LEADER:
        DOTS = 'dots'
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class TOCGenerator:
    """目录生成器"""
    
    def __init__(self):
        # 标题级别识别模式
        self.heading_patterns = {
            1: [
                r'^第[一二三四五六七八九十\d]+[章节]',
                r'^Chapter\s+\d+',
                r'^\d+\s+\S',
                r'^\d+\.\s+\S'
            ],
            2: [
                r'^\d+\.\d+\s+\S',
                r'^[一二三四五六七八九十]+、',
            ],
            3: [
                r'^\d+\.\d+\.\d+\s+\S',
                r'^（[一二三四五六七八九十\d]+）',
                r'^\([一二三四五六七八九十\d]+\)',
            ]
        }
        
        # 特殊标题（不编号但要包含在目录中）
        self.special_headings = [
            '摘要', 'Abstract', '目录', '参考文献', 
            '致谢', '附录', '承诺书', '诚信承诺书'
        ]
    
    def generate_toc(self, doc, structure):
        """
        生成目录
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        # 查找或创建目录位置
        toc_idx = structure.get('toc', -1)
        
        if toc_idx < 0:
            # 如果没有目录，在摘要后创建
            toc_idx = self._find_toc_position(doc, structure)
        
        # 收集标题信息
        headings = self._collect_headings(doc, structure)
        
        # 生成目录内容
        self._insert_toc(doc, toc_idx, headings)
    
    def update_toc(self, doc, structure):
        """
        更新现有目录
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        toc_idx = structure.get('toc', -1)
        
        if toc_idx >= 0:
            # 清除现有目录内容
            toc_end = self._find_toc_end(doc, toc_idx)
            
            # 收集新的标题信息
            headings = self._collect_headings(doc, structure)
            
            # 删除旧目录内容
            self._clear_toc_content(doc, toc_idx, toc_end)
            
            # 插入新目录
            self._insert_toc_content(doc, toc_idx, headings)
    
    def _find_toc_position(self, doc, structure):
        """找到目录应该插入的位置"""
        # 通常在摘要之后，正文之前
        abstract_en = structure.get('abstract_en', -1)
        main_start = structure.get('main_start', -1)
        
        if abstract_en >= 0 and main_start > abstract_en:
            # 在英文摘要后插入
            return abstract_en + 3  # 留出一些空间
        elif main_start > 0:
            # 在正文前插入
            return main_start - 1
        else:
            # 默认在第5个段落
            return min(5, len(doc.paragraphs))
    
    def _collect_headings(self, doc, structure):
        """收集所有标题及其信息"""
        headings = []
        
        # 添加特殊部分（前置部分）
        if structure.get('abstract_cn', -1) >= 0:
            headings.append({
                'text': '摘要',
                'level': 0,
                'page': 'I',  # 罗马数字
                'type': 'special'
            })
        
        if structure.get('abstract_en', -1) >= 0:
            headings.append({
                'text': 'Abstract',
                'level': 0,
                'page': 'II',
                'type': 'special'
            })
        
        # 收集正文标题
        main_start = structure.get('main_start', 0)
        chapter_num = 0
        current_page = 1  # 正文从第1页开始
        
        for i in range(main_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            
            # 检测标题级别
            level = self._detect_heading_level(para)
            
            if level > 0:
                # 估算页码（简化处理）
                page_num = self._estimate_page_number(i, main_start, len(doc.paragraphs))
                
                headings.append({
                    'text': para.text.strip(),
                    'level': level,
                    'page': str(page_num),
                    'type': 'normal'
                })
            
            # 检测特殊标题
            elif self._is_special_heading(para):
                headings.append({
                    'text': para.text.strip(),
                    'level': 0,
                    'page': str(self._estimate_page_number(i, main_start, len(doc.paragraphs))),
                    'type': 'special'
                })
        
        return headings
    
    def _detect_heading_level(self, paragraph):
        """检测标题级别"""
        text = paragraph.text.strip()
        
        # 检查是否为粗体（通常是标题）
        is_bold = paragraph.runs and paragraph.runs[0].font.bold
        
        # 按级别检查
        for level, patterns in self.heading_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text):
                    return level
        
        # 如果是粗体且较短，可能是标题
        if is_bold and len(text) < 50:
            # 根据缩进判断级别
            indent = paragraph.paragraph_format.left_indent
            if indent is None:
                return 1
            elif indent < Pt(24):
                return 2
            else:
                return 3
        
        return 0
    
    def _is_special_heading(self, paragraph):
        """检查是否为特殊标题"""
        text = paragraph.text.strip()
        
        # 检查是否匹配特殊标题
        for heading in self.special_headings:
            if heading in text and len(text) < 20:
                # 检查格式特征
                if paragraph.runs and paragraph.runs[0].font.bold:
                    return True
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return True
        
        return False
    
    def _estimate_page_number(self, para_index, main_start, total_paras):
        """估算页码（简化算法）"""
        # 假设平均每页30个段落
        paras_per_page = 30
        
        if para_index < main_start:
            # 前置部分，返回罗马数字
            page = (para_index // paras_per_page) + 1
            return self._arabic_to_roman(page)
        else:
            # 正文部分，返回阿拉伯数字
            page = ((para_index - main_start) // paras_per_page) + 1
            return page
    
    def _arabic_to_roman(self, num):
        """阿拉伯数字转罗马数字"""
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4,
            1
        ]
        syms = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV",
            "I"
        ]
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num
    
    def _find_toc_end(self, doc, toc_start):
        """找到目录结束位置"""
        # 从目录标题开始，找到下一个主要标题或正文开始
        for i in range(toc_start + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            
            # 检查是否为新的主要部分
            if self._is_special_heading(para) or self._detect_heading_level(para) == 1:
                return i
            
            # 检查是否有分页符
            for run in para.runs:
                if run._element.xml.find('w:br') != -1:
                    return i
        
        # 默认返回目录后10个段落
        return min(toc_start + 10, len(doc.paragraphs))
    
    def _clear_toc_content(self, doc, toc_start, toc_end):
        """清除目录内容（保留标题）"""
        # 保留目录标题，删除内容
        for i in range(toc_end - 1, toc_start, -1):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                # 不删除目录标题
                if '目录' not in para.text:
                    para._element.getparent().remove(para._element)
    
    def _insert_toc(self, doc, position, headings):
        """在指定位置插入完整目录"""
        # 如果位置处没有目录标题，添加一个
        if position >= len(doc.paragraphs) or '目录' not in doc.paragraphs[position].text:
            # 插入目录标题
            toc_title = doc.add_paragraph()
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = toc_title.add_run('目录')
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(18)  # 二号
            run.font.bold = True
            
            # 移动到正确位置
            if position < len(doc.paragraphs):
                doc.paragraphs[position]._element.addprevious(toc_title._element)
        
        # 插入目录内容
        self._insert_toc_content(doc, position, headings)
    
    def _insert_toc_content(self, doc, toc_idx, headings):
        """插入目录内容"""
        # 在目录标题后插入内容
        insert_point = toc_idx + 1
        
        for heading in headings:
            # 创建目录项段落
            if insert_point < len(doc.paragraphs):
                para = doc.paragraphs[insert_point].insert_paragraph_before()
            else:
                para = doc.add_paragraph()
            
            # 设置缩进
            if heading['level'] == 2:
                para.paragraph_format.left_indent = Pt(24)
            elif heading['level'] == 3:
                para.paragraph_format.left_indent = Pt(48)
            
            # 设置段落格式
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(22)
            
            # 添加制表位（右对齐页码）
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.clear_all()
            # 简化制表位设置
            try:
                tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT)
            except:
                # 兼容性处理
                pass
            
            # 添加标题文本
            run_title = para.add_run(heading['text'])
            run_title.font.name = '宋体'
            run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_title.font.size = Pt(12)  # 小四号
            
            # 特殊标题加粗
            if heading['type'] == 'special':
                run_title.font.bold = True
            
            # 添加制表符
            para.add_run('\t')
            
            # 添加页码
            run_page = para.add_run(heading['page'])
            run_page.font.name = '宋体'
            run_page._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_page.font.size = Pt(12)
    
    def add_toc_field(self, doc, position):
        """
        添加Word的目录域（高级功能）
        注：这会创建一个可以在Word中更新的动态目录
        """
        para = doc.add_paragraph()
        
        # 创建目录域
        run = para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '  # 目录域代码
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # 移动到指定位置
        if position < len(doc.paragraphs):
            doc.paragraphs[position]._element.addprevious(para._element)


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加一些测试内容
    # 摘要
    p1 = doc.add_paragraph("摘要")
    p1.runs[0].font.bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("这是中文摘要的内容...")
    
    # Abstract
    p2 = doc.add_paragraph("Abstract")
    p2.runs[0].font.bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("This is the abstract content...")
    
    # 目录（占位）
    p_toc = doc.add_paragraph("目录")
    p_toc.runs[0].font.bold = True
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 正文
    doc.add_page_break()
    
    # 第一章
    p3 = doc.add_paragraph("第一章 绪论")
    p3.runs[0].font.bold = True
    
    p4 = doc.add_paragraph("1.1 研究背景")
    p4.runs[0].font.bold = True
    
    doc.add_paragraph("1.1.1 国内外研究现状")
    
    # 第二章
    p5 = doc.add_paragraph("第二章 相关技术")
    p5.runs[0].font.bold = True
    
    p6 = doc.add_paragraph("2.1 技术概述")
    p6.runs[0].font.bold = True
    
    # 参考文献
    p7 = doc.add_paragraph("参考文献")
    p7.runs[0].font.bold = True
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 测试目录生成
    generator = TOCGenerator()
    structure = {
        'abstract_cn': 0,
        'abstract_en': 2,
        'toc': 4,
        'main_start': 5
    }
    
    generator.update_toc(doc, structure)
    
    # 保存测试文档
    doc.save('test_toc.docx')
    print("测试文档已生成：test_toc.docx")
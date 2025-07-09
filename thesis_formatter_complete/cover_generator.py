#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
封面生成器模块
生成符合江西财经大学现代经济管理学院规范的论文封面
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START as WD_SECTION
import datetime

class CoverGenerator:
    """论文封面生成器"""
    
    def __init__(self):
        self.default_info = {
            'title': '基于深度学习的图像识别研究',
            'major': '计算机科学与技术',
            'class': '计科1901',
            'student_id': '20190001',
            'name': '张三',
            'advisor': '李教授',
            'date': datetime.datetime.now().strftime('%Y年%m月')
        }
    
    def generate_cover(self, doc, thesis_info=None):
        """
        生成论文封面
        
        Args:
            doc: Document对象
            thesis_info: 论文信息字典，包含title、major、class等
        """
        if thesis_info is None:
            thesis_info = self.default_info
        
        # 如果文档不是空的，在最前面插入分页符
        if len(doc.paragraphs) > 0:
            # 在文档开头插入一个段落
            first_para = doc.paragraphs[0]
            new_para = first_para.insert_paragraph_before()
            # 添加分页符而不是分节符
            new_para.add_run().add_break()
        
        # 封面从新的一页开始
        cover_paras = []
        
        # 1. 学校名称
        p_school = doc.add_paragraph()
        p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_school = p_school.add_run("江西财经大学现代经济管理学院")
        self._set_font(run_school, '宋体', Pt(22), bold=True)
        cover_paras.append(p_school)
        
        # 2. 空行
        p_space1 = doc.add_paragraph()
        p_space1.paragraph_format.space_after = Pt(36)
        cover_paras.append(p_space1)
        
        # 3. "毕业论文"标题
        p_type = doc.add_paragraph()
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_type = p_type.add_run("毕业论文")
        self._set_font(run_type, '宋体', Pt(26), bold=True)
        cover_paras.append(p_type)
        
        # 4. 空行
        p_space2 = doc.add_paragraph()
        p_space2.paragraph_format.space_after = Pt(48)
        cover_paras.append(p_space2)
        
        # 5. 论文题目
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = thesis_info.get('title', self.default_info['title'])
        
        # 如果标题太长，需要换行
        if len(title_text) > 20:
            # 找到合适的断点
            mid_point = len(title_text) // 2
            # 在中点附近找到合适的断词位置
            for i in range(mid_point - 5, mid_point + 5):
                if i < len(title_text) and title_text[i] in '的与和及或':
                    mid_point = i + 1
                    break
            
            run_title1 = p_title.add_run(title_text[:mid_point])
            self._set_font(run_title1, '宋体', Pt(16), bold=True)
            p_title.add_run('\n')
            run_title2 = p_title.add_run(title_text[mid_point:])
            self._set_font(run_title2, '宋体', Pt(16), bold=True)
        else:
            run_title = p_title.add_run(title_text)
            self._set_font(run_title, '宋体', Pt(16), bold=True)
        
        p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        cover_paras.append(p_title)
        
        # 6. 空行
        p_space3 = doc.add_paragraph()
        p_space3.paragraph_format.space_after = Pt(72)
        cover_paras.append(p_space3)
        
        # 7. 学生信息
        info_items = [
            ('专　　业：', thesis_info.get('major', self.default_info['major'])),
            ('班　　级：', thesis_info.get('class', self.default_info['class'])),
            ('学　　号：', thesis_info.get('student_id', self.default_info['student_id'])),
            ('姓　　名：', thesis_info.get('name', self.default_info['name'])),
            ('指导教师：', thesis_info.get('advisor', self.default_info['advisor']))
        ]
        
        for label, value in info_items:
            p_info = doc.add_paragraph()
            p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 使用表格对齐，确保冒号对齐
            run_label = p_info.add_run(label)
            self._set_font(run_label, '宋体', Pt(16))
            
            run_value = p_info.add_run(value)
            self._set_font(run_value, '宋体', Pt(16))
            run_value.underline = True
            
            p_info.paragraph_format.space_after = Pt(18)
            cover_paras.append(p_info)
        
        # 8. 空行
        p_space4 = doc.add_paragraph()
        p_space4.paragraph_format.space_after = Pt(72)
        cover_paras.append(p_space4)
        
        # 9. 日期
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = p_date.add_run(thesis_info.get('date', self.default_info['date']))
        self._set_font(run_date, '宋体', Pt(14))
        cover_paras.append(p_date)
        
        # 10. 在封面后插入分页符
        doc.add_page_break()
        
        return cover_paras
    
    def _set_font(self, run, font_name, size, bold=False):
        """设置字体格式"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size
        run.font.bold = bold
    
    def validate_thesis_info(self, thesis_info):
        """验证论文信息的完整性"""
        required_fields = ['title', 'major', 'class', 'student_id', 'name', 'advisor']
        missing_fields = []
        
        for field in required_fields:
            if not thesis_info.get(field):
                missing_fields.append(field)
        
        if missing_fields:
            return False, f"缺少必需字段：{', '.join(missing_fields)}"
        
        # 验证标题长度
        if len(thesis_info['title']) > 40:
            return False, "论文标题不能超过40个字符"
        
        return True, "信息完整"


class CommitmentGenerator:
    """诚信承诺书生成器"""
    
    def generate_commitment(self, doc):
        """生成诚信承诺书页"""
        # 标题
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("诚信承诺书")
        self._set_font(run_title, '宋体', Pt(18), bold=True)
        p_title.paragraph_format.space_after = Pt(36)
        
        # 承诺内容
        commitment_text = """本人郑重承诺：

　　所呈交的毕业论文是本人在导师指导下独立进行研究工作所取得的研究成果。除了文中特别加以标注和致谢的地方外，论文中不包含其他人或集体已经发表或撰写过的研究成果，也不包含为获得江西财经大学现代经济管理学院或其它教育机构的学位或证书而使用过的材料。

　　本人承诺，如违反上述声明，愿意承担由此引发的一切责任和后果。"""
        
        # 分段添加承诺内容
        for paragraph in commitment_text.strip().split('\n\n'):
            p_content = doc.add_paragraph()
            if paragraph.startswith('　　'):
                # 正文段落
                p_content.paragraph_format.first_line_indent = Pt(28)
                p_content.paragraph_format.line_spacing = Pt(28)
            else:
                # 标题行
                p_content.paragraph_format.space_after = Pt(12)
            
            run_content = p_content.add_run(paragraph.strip())
            self._set_font(run_content, '宋体', Pt(14))
        
        # 空行
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(72)
        
        # 签名区域
        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sign.paragraph_format.right_indent = Cm(3)
        run_sign = p_sign.add_run("承诺人（签名）：" + "_" * 20)
        self._set_font(run_sign, '宋体', Pt(14))
        
        # 日期
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.right_indent = Cm(3)
        p_date.paragraph_format.space_before = Pt(12)
        run_date = p_date.add_run("日期：____年____月____日")
        self._set_font(run_date, '宋体', Pt(14))
        
        # 添加分页符
        doc.add_page_break()
    
    def _set_font(self, run, font_name, size, bold=False):
        """设置字体格式"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size
        run.font.bold = bold


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 测试封面生成
    cover_gen = CoverGenerator()
    thesis_info = {
        'title': '基于Python的毕业论文格式化工具设计与实现',
        'major': '软件工程',
        'class': '软工2001',
        'student_id': '20200101',
        'name': '王小明',
        'advisor': '张教授',
        'date': '2024年5月'
    }
    
    cover_gen.generate_cover(doc, thesis_info)
    
    # 测试诚信承诺书生成
    commitment_gen = CommitmentGenerator()
    commitment_gen.generate_commitment(doc)
    
    # 保存测试文档
    doc.save('test_cover_commitment.docx')
    print("测试文档已生成：test_cover_commitment.docx")
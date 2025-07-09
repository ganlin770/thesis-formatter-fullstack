#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
页眉处理器模块
负责添加和管理文档页眉，确保符合江西财经大学现代经济管理学院要求
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import logging
from typing import Dict, List, Optional

class HeaderHandler:
    """页眉处理器：管理文档页眉设置"""
    
    # 页眉配置
    HEADER_CONFIG = {
        'text': '江西财经大学现代经济管理学院普通本科毕业论文',
        'font_name': '宋体',
        'font_size': 10.5,  # 五号
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'distance_from_top': Cm(1.5),  # 距顶端1.5厘米
        'border_bottom': True  # 页眉下方加细线
    }
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def add_header_to_section(self, section) -> None:
        """
        为指定节添加页眉
        
        Args:
            section: docx section对象
        """
        try:
            # 获取页眉
            header = section.header
            
            # 清空现有页眉内容
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.clear()
            
            # 设置页眉文本
            run = header_para.add_run(self.HEADER_CONFIG['text'])
            
            # 设置字体和字号
            run.font.name = self.HEADER_CONFIG['font_name']
            run.font.size = Pt(self.HEADER_CONFIG['font_size'])
            
            # 设置段落对齐
            header_para.alignment = self.HEADER_CONFIG['alignment']
            
            # 设置页眉距离顶端的距离
            section.top_margin = self.HEADER_CONFIG['distance_from_top']
            
            # 添加页眉下方细线
            if self.HEADER_CONFIG['border_bottom']:
                self._add_header_border(header_para)
            
            self.logger.info(f"已为节添加页眉: {self.HEADER_CONFIG['text']}")
            
        except Exception as e:
            self.logger.error(f"添加页眉失败: {e}")
    
    def _add_header_border(self, paragraph) -> None:
        """
        为页眉段落添加下边框
        
        Args:
            paragraph: 页眉段落对象
        """
        try:
            # 获取段落边框设置
            pPr = paragraph._element.get_or_add_pPr()
            
            # 创建边框元素
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            border_xml = f'''
            <w:pBdr {nsdecls('w')}>
                <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>
            </w:pBdr>
            '''
            
            border_element = parse_xml(border_xml)
            pPr.append(border_element)
            
        except Exception as e:
            self.logger.warning(f"添加页眉边框失败: {e}")
    
    def configure_header_sections(self, document: Document, structure: Dict) -> None:
        """
        配置页眉显示范围（仅正文部分）
        
        Args:
            document: Word文档对象
            structure: 文档结构信息
        """
        try:
            # 获取正文开始位置
            main_start = structure.get('main_start', -1)
            
            if main_start == -1:
                self.logger.warning("未找到正文开始位置，将为所有节添加页眉")
                # 如果找不到正文开始，为所有节添加页眉
                for section in document.sections:
                    self.add_header_to_section(section)
                return
            
            # 分析文档节结构
            sections_info = self._analyze_sections(document, structure)
            
            # 只为正文部分的节添加页眉
            for i, section in enumerate(document.sections):
                section_info = sections_info.get(i, {})
                
                if section_info.get('is_main_content', False):
                    self.add_header_to_section(section)
                    self.logger.info(f"为第{i+1}节添加页眉（正文区域）")
                else:
                    self._remove_header_from_section(section)
                    self.logger.info(f"第{i+1}节不添加页眉（非正文区域）")
            
        except Exception as e:
            self.logger.error(f"配置页眉节失败: {e}")
    
    def _analyze_sections(self, document: Document, structure: Dict) -> Dict:
        """
        分析文档节结构
        
        Args:
            document: Word文档对象
            structure: 文档结构信息
            
        Returns:
            节结构信息字典
        """
        sections_info = {}
        
        # 获取关键位置
        abstract_cn = structure.get('abstract_cn', 0)
        main_start = structure.get('main_start', -1)
        references = structure.get('references', len(document.paragraphs))
        
        # 估算每个节包含的段落范围
        paragraphs_per_section = len(document.paragraphs) // len(document.sections)
        
        for i, section in enumerate(document.sections):
            section_start = i * paragraphs_per_section
            section_end = (i + 1) * paragraphs_per_section
            
            # 判断这个节是否包含正文内容
            is_main_content = (
                main_start != -1 and 
                section_start <= main_start < section_end or
                main_start < section_start < references
            )
            
            sections_info[i] = {
                'section_start': section_start,
                'section_end': section_end,
                'is_main_content': is_main_content
            }
        
        return sections_info
    
    def _remove_header_from_section(self, section) -> None:
        """
        从指定节移除页眉
        
        Args:
            section: docx section对象
        """
        try:
            header = section.header
            if header.paragraphs:
                header_para = header.paragraphs[0]
                header_para.clear()
                
        except Exception as e:
            self.logger.warning(f"移除页眉失败: {e}")
    
    def validate_headers(self, document: Document) -> Dict[str, List[str]]:
        """
        验证页眉设置
        
        Args:
            document: Word文档对象
            
        Returns:
            包含错误和警告的字典
        """
        errors = []
        warnings = []
        
        for i, section in enumerate(document.sections):
            try:
                header = section.header
                
                if not header.paragraphs:
                    warnings.append(f"第{i+1}节：页眉为空")
                    continue
                
                header_para = header.paragraphs[0]
                header_text = header_para.text.strip()
                
                # 检查页眉文本
                if header_text and header_text != self.HEADER_CONFIG['text']:
                    warnings.append(f"第{i+1}节：页眉文本不符合要求")
                
                # 检查字体设置
                if header_para.runs:
                    run = header_para.runs[0]
                    
                    if run.font.name != self.HEADER_CONFIG['font_name']:
                        errors.append(f"第{i+1}节：页眉字体应为{self.HEADER_CONFIG['font_name']}")
                    
                    expected_size = Pt(self.HEADER_CONFIG['font_size'])
                    if run.font.size != expected_size:
                        errors.append(f"第{i+1}节：页眉字号应为{self.HEADER_CONFIG['font_size']}磅")
                
                # 检查对齐方式
                if header_para.alignment != self.HEADER_CONFIG['alignment']:
                    errors.append(f"第{i+1}节：页眉应该居中对齐")
                    
            except Exception as e:
                errors.append(f"第{i+1}节：页眉验证失败 - {e}")
        
        return {'errors': errors, 'warnings': warnings}
    
    def add_page_header_to_document(self, document: Document, structure: Dict) -> bool:
        """
        为整个文档添加页眉（主要接口）
        
        Args:
            document: Word文档对象
            structure: 文档结构信息
            
        Returns:
            是否成功添加页眉
        """
        try:
            self.configure_header_sections(document, structure)
            
            # 验证页眉设置
            validation_result = self.validate_headers(document)
            
            if validation_result['errors']:
                self.logger.warning(f"页眉设置存在错误: {validation_result['errors']}")
            
            self.logger.info("页眉添加完成")
            return True
            
        except Exception as e:
            self.logger.error(f"添加页眉失败: {e}")
            return False
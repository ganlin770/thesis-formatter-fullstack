#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
图表编号处理模块
实现按章节编号的图表标题格式化（图1.1、表1.1）
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

class FigureTableHandler:
    """图表编号处理器"""
    
    def __init__(self):
        self.figure_counters = {}  # 章节号 -> 图编号
        self.table_counters = {}   # 章节号 -> 表编号
        self.current_chapter = 1
        
        # 图表标题识别模式
        self.figure_patterns = [
            r'^图\s*[\d.]+',
            r'^图\s*\d+[-_]\d+',
            r'^Fig\.?\s*\d+',
            r'^Figure\s*\d+'
        ]
        
        self.table_patterns = [
            r'^表\s*[\d.]+',
            r'^表\s*\d+[-_]\d+',
            r'^Tab\.?\s*\d+',
            r'^Table\s*\d+'
        ]
        
        # 章节标题模式
        self.chapter_patterns = [
            r'^第([一二三四五六七八九十\d]+)[章节]',
            r'^(\d+)[.\s]+\S',
            r'^Chapter\s+(\d+)'
        ]
    
    def process_figures_and_tables(self, doc, structure):
        """
        处理整个文档的图表编号
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        main_start = structure.get('main_start', 0)
        if main_start < 0:
            main_start = 0
        
        # 遍历文档，处理章节和图表
        for i in range(main_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            
            # 检测章节
            chapter_num = self._detect_chapter(para)
            if chapter_num:
                self.current_chapter = chapter_num
                # 重置该章节的图表计数器
                if chapter_num not in self.figure_counters:
                    self.figure_counters[chapter_num] = 0
                if chapter_num not in self.table_counters:
                    self.table_counters[chapter_num] = 0
            
            # 检测并处理图标题
            if self._is_figure_caption(para):
                self._format_figure_caption(para)
            
            # 检测并处理表标题
            elif self._is_table_caption(para):
                self._format_table_caption(para)
    
    def _detect_chapter(self, paragraph):
        """检测章节编号"""
        text = paragraph.text.strip()
        
        for pattern in self.chapter_patterns:
            match = re.match(pattern, text)
            if match:
                chapter_str = match.group(1)
                # 转换中文数字
                if chapter_str in '一二三四五六七八九十':
                    return self._chinese_to_arabic(chapter_str)
                else:
                    return int(chapter_str)
        
        return None
    
    def _chinese_to_arabic(self, chinese_num):
        """中文数字转阿拉伯数字"""
        num_map = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10
        }
        return num_map.get(chinese_num, 1)
    
    def _is_figure_caption(self, paragraph):
        """判断是否为图标题"""
        text = paragraph.text.strip()
        
        # 检查是否匹配图标题模式
        for pattern in self.figure_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        # 检查是否包含"图"且在合理长度内
        if '图' in text and len(text) < 100:
            # 检查是否在表格或其他元素附近
            return self._is_likely_caption(paragraph)
        
        return False
    
    def _is_table_caption(self, paragraph):
        """判断是否为表标题"""
        text = paragraph.text.strip()
        
        # 检查是否匹配表标题模式
        for pattern in self.table_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        # 检查是否包含"表"且在合理长度内
        if '表' in text and len(text) < 100:
            return self._is_likely_caption(paragraph)
        
        return False
    
    def _is_likely_caption(self, paragraph):
        """判断段落是否可能是标题"""
        # 检查格式特征
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        
        # 检查字体特征
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # 小四号字体
            if first_run.font.size and first_run.font.size <= Pt(12):
                return True
        
        return False
    
    def _format_figure_caption(self, paragraph):
        """格式化图标题"""
        # 增加当前章节的图计数
        if self.current_chapter not in self.figure_counters:
            self.figure_counters[self.current_chapter] = 0
        self.figure_counters[self.current_chapter] += 1
        
        # 生成新的图编号
        new_number = f"图{self.current_chapter}.{self.figure_counters[self.current_chapter]}"
        
        # 替换原有编号
        original_text = paragraph.text.strip()
        new_text = self._replace_figure_number(original_text, new_number)
        
        # 清空段落并重新设置
        paragraph.clear()
        run = paragraph.add_run(new_text)
        
        # 设置格式
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)  # 小四号
        
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = Pt(22)
    
    def _format_table_caption(self, paragraph):
        """格式化表标题"""
        # 增加当前章节的表计数
        if self.current_chapter not in self.table_counters:
            self.table_counters[self.current_chapter] = 0
        self.table_counters[self.current_chapter] += 1
        
        # 生成新的表编号
        new_number = f"表{self.current_chapter}.{self.table_counters[self.current_chapter]}"
        
        # 替换原有编号
        original_text = paragraph.text.strip()
        new_text = self._replace_table_number(original_text, new_number)
        
        # 清空段落并重新设置
        paragraph.clear()
        run = paragraph.add_run(new_text)
        
        # 设置格式
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)  # 小四号
        
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = Pt(22)
    
    def _replace_figure_number(self, text, new_number):
        """替换图编号"""
        # 尝试多种替换模式
        patterns = [
            r'图\s*[\d.]+\s*',
            r'图\s*\d+[-_]\d+\s*',
            r'Fig\.?\s*\d+\s*',
            r'Figure\s*\d+\s*'
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                # 保留编号后的内容
                remaining = re.sub(pattern, '', text, count=1, flags=re.IGNORECASE)
                return f"{new_number} {remaining}".strip()
        
        # 如果没有匹配到编号，在"图"后添加编号
        if '图' in text:
            return text.replace('图', new_number, 1)
        
        # 默认在开头添加编号
        return f"{new_number} {text}"
    
    def _replace_table_number(self, text, new_number):
        """替换表编号"""
        # 尝试多种替换模式
        patterns = [
            r'表\s*[\d.]+\s*',
            r'表\s*\d+[-_]\d+\s*',
            r'Tab\.?\s*\d+\s*',
            r'Table\s*\d+\s*'
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                # 保留编号后的内容
                remaining = re.sub(pattern, '', text, count=1, flags=re.IGNORECASE)
                return f"{new_number} {remaining}".strip()
        
        # 如果没有匹配到编号，在"表"后添加编号
        if '表' in text:
            return text.replace('表', new_number, 1)
        
        # 默认在开头添加编号
        return f"{new_number} {text}"
    
    def update_cross_references(self, doc):
        """更新文档中的交叉引用"""
        # 存储所有图表编号
        figure_refs = {}  # "图1-1" -> "图1.1"
        table_refs = {}   # "表1-1" -> "表1.1"
        
        # 第一遍：收集所有图表编号
        for para in doc.paragraphs:
            text = para.text
            
            # 收集图编号
            figure_matches = re.findall(r'图\s*(\d+)[-_](\d+)', text)
            for match in figure_matches:
                old_ref = f"图{match[0]}-{match[1]}"
                new_ref = f"图{match[0]}.{match[1]}"
                figure_refs[old_ref] = new_ref
            
            # 收集表编号
            table_matches = re.findall(r'表\s*(\d+)[-_](\d+)', text)
            for match in table_matches:
                old_ref = f"表{match[0]}-{match[1]}"
                new_ref = f"表{match[0]}.{match[1]}"
                table_refs[old_ref] = new_ref
        
        # 第二遍：替换所有引用
        for para in doc.paragraphs:
            if self._is_figure_caption(para) or self._is_table_caption(para):
                continue  # 跳过标题本身
            
            text = para.text
            changed = False
            
            # 替换图引用
            for old_ref, new_ref in figure_refs.items():
                if old_ref in text:
                    text = text.replace(old_ref, new_ref)
                    changed = True
            
            # 替换表引用
            for old_ref, new_ref in table_refs.items():
                if old_ref in text:
                    text = text.replace(old_ref, new_ref)
                    changed = True
            
            # 如果有改变，更新段落
            if changed:
                # 保存原有格式
                if para.runs:
                    font_name = para.runs[0].font.name
                    font_size = para.runs[0].font.size
                    font_bold = para.runs[0].font.bold
                    
                    para.clear()
                    run = para.add_run(text)
                    run.font.name = font_name
                    if font_name:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = font_size
                    run.font.bold = font_bold


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加第一章
    p1 = doc.add_paragraph("第一章 绪论")
    p1.runs[0].font.bold = True
    
    doc.add_paragraph("这是第一章的内容...")
    
    # 添加图
    p_fig1 = doc.add_paragraph("图1-1 系统架构图")
    p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("如图1-1所示，系统架构包含...")
    
    # 添加表
    p_tab1 = doc.add_paragraph("表 1 实验结果对比")
    p_tab1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加第二章
    p2 = doc.add_paragraph("第二章 相关工作")
    p2.runs[0].font.bold = True
    
    # 添加更多图表
    p_fig2 = doc.add_paragraph("图 算法流程")
    p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_tab2 = doc.add_paragraph("表2-1 参数设置")
    p_tab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 测试处理器
    handler = FigureTableHandler()
    structure = {'main_start': 0}
    
    handler.process_figures_and_tables(doc, structure)
    handler.update_cross_references(doc)
    
    # 保存测试文档
    doc.save('test_figure_table.docx')
    print("测试文档已生成：test_figure_table.docx")
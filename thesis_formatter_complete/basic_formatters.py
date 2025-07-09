#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基础格式化器模块
提供标题、段落、摘要和参考文献的基础格式化功能
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


class HeadingFormatter:
    """标题格式化器"""
    
    def __init__(self, document, config):
        self.document = document
        self.config = config
    
    def format_headings(self, structure):
        """格式化所有标题"""
        for i, para in enumerate(self.document.paragraphs):
            level = self._detect_heading_level(para)
            if level > 0:
                self._format_heading(para, level)
    
    def _detect_heading_level(self, paragraph):
        """检测标题级别"""
        text = paragraph.text.strip()
        
        # 一级标题模式
        if re.match(r'^第[一二三四五六七八九十\d]+[章节]', text):
            return 1
        # 二级标题模式
        elif re.match(r'^\d+\.\d+\s+', text):
            return 2
        # 三级标题模式
        elif re.match(r'^\d+\.\d+\.\d+\s+', text):
            return 3
        
        return 0
    
    def _format_heading(self, paragraph, level):
        """格式化标题"""
        # 保存原文本
        text = paragraph.text
        
        # 清空并重新设置
        paragraph.clear()
        run = paragraph.add_run(text)
        
        # 设置字体
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.bold = True
        
        # 根据级别设置字号
        if level == 1:
            run.font.size = Pt(18)  # 二号
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(16)  # 三号
        else:
            run.font.size = Pt(14)  # 四号


class ParagraphFormatter:
    """段落格式化器"""
    
    def __init__(self, document, config):
        self.document = document
        self.config = config
    
    def format_paragraphs(self, structure):
        """格式化正文段落"""
        main_start = structure.get('main_start', 0)
        ref_start = structure.get('references', len(self.document.paragraphs))
        
        for i in range(main_start, min(ref_start, len(self.document.paragraphs))):
            para = self.document.paragraphs[i]
            
            # 跳过标题
            if self._is_heading(para):
                continue
            
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 格式化段落
            self._format_paragraph(para)
    
    def _is_heading(self, paragraph):
        """判断是否为标题"""
        if paragraph.runs and paragraph.runs[0].font.bold:
            return True
        return False
    
    def _format_paragraph(self, paragraph):
        """格式化单个段落"""
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        # 设置字体
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)  # 小四号


class AbstractFormatter:
    """摘要格式化器"""
    
    def __init__(self, document, config):
        self.document = document
        self.config = config
    
    def format_abstract(self, structure):
        """格式化中英文摘要"""
        # 格式化中文摘要
        cn_idx = structure.get('abstract_cn', -1)
        if cn_idx >= 0:
            self._format_cn_abstract(cn_idx)
        
        # 格式化英文摘要
        en_idx = structure.get('abstract_en', -1)
        if en_idx >= 0:
            self._format_en_abstract(en_idx)
    
    def _format_cn_abstract(self, start_idx):
        """格式化中文摘要"""
        # 格式化标题
        if start_idx < len(self.document.paragraphs):
            title_para = self.document.paragraphs[start_idx]
            self._format_abstract_title(title_para, '摘要')
        
        # 格式化内容
        for i in range(start_idx + 1, min(start_idx + 10, len(self.document.paragraphs))):
            para = self.document.paragraphs[i]
            
            # 遇到关键词或其他标题停止
            if '关键词' in para.text or 'Abstract' in para.text:
                break
            
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 格式化摘要内容
            self._format_abstract_content(para, 'cn')
    
    def _format_en_abstract(self, start_idx):
        """格式化英文摘要"""
        # 格式化标题
        if start_idx < len(self.document.paragraphs):
            title_para = self.document.paragraphs[start_idx]
            self._format_abstract_title(title_para, 'Abstract')
        
        # 格式化内容
        for i in range(start_idx + 1, min(start_idx + 10, len(self.document.paragraphs))):
            para = self.document.paragraphs[i]
            
            # 遇到关键词或其他标题停止
            if 'Keywords' in para.text or '第' in para.text:
                break
            
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 格式化摘要内容
            self._format_abstract_content(para, 'en')
    
    def _format_abstract_title(self, paragraph, title):
        """格式化摘要标题"""
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run(title)
        run.font.name = '宋体' if title == '摘要' else 'Times New Roman'
        if title == '摘要':
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(18)  # 二号
        run.font.bold = True
    
    def _format_abstract_content(self, paragraph, lang='cn'):
        """格式化摘要内容"""
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = Pt(22)
        
        # 设置字体
        for run in paragraph.runs:
            if lang == 'cn':
                run.font.name = '楷体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            else:
                run.font.name = 'Times New Roman'
            run.font.size = Pt(12)  # 小四号


class ReferenceFormatter:
    """参考文献格式化器"""
    
    def __init__(self, document, config):
        self.document = document
        self.config = config
    
    def format_references(self, structure):
        """格式化参考文献"""
        ref_idx = structure.get('references', -1)
        if ref_idx < 0:
            return
        
        # 格式化标题
        if ref_idx < len(self.document.paragraphs):
            title_para = self.document.paragraphs[ref_idx]
            self._format_reference_title(title_para)
        
        # 格式化参考文献条目
        for i in range(ref_idx + 1, len(self.document.paragraphs)):
            para = self.document.paragraphs[i]
            
            # 遇到其他标题停止
            if self._is_new_section(para):
                break
            
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 格式化参考文献
            self._format_reference_item(para)
    
    def _format_reference_title(self, paragraph):
        """格式化参考文献标题"""
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run('参考文献')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(18)  # 二号
        run.font.bold = True
    
    def _format_reference_item(self, paragraph):
        """格式化参考文献条目"""
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = -Pt(21)  # 悬挂缩进
        paragraph.paragraph_format.line_spacing = Pt(18)  # 18磅行距
        
        # 设置字体
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)  # 五号
    
    def _is_new_section(self, paragraph):
        """判断是否为新章节"""
        text = paragraph.text.strip()
        if any(keyword in text for keyword in ['致谢', '附录', 'Appendix']):
            if paragraph.runs and paragraph.runs[0].font.bold:
                return True
        return False
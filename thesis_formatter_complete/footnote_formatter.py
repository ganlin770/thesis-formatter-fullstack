#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
脚注格式化模块
实现小五号宋体，行距12pt，每页重新编号的脚注格式
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class FootnoteFormatter:
    """脚注格式化器"""
    
    def __init__(self):
        self.footnote_pattern = re.compile(r'\[\d+\]|\(\d+\)|[\u2460-\u2473]|①②③④⑤⑥⑦⑧⑨⑩')
        
    def format_footnotes(self, doc):
        """
        格式化文档中的脚注
        
        Args:
            doc: Document对象
        """
        # 查找并格式化脚注引用
        self._format_footnote_references(doc)
        
        # 尝试格式化脚注内容（如果能访问）
        self._format_footnote_content(doc)
    
    def _format_footnote_references(self, doc):
        """格式化正文中的脚注引用标记"""
        for para in doc.paragraphs:
            for run in para.runs:
                # 检查是否包含脚注标记
                if self._contains_footnote_marker(run.text):
                    # 格式化脚注引用
                    self._format_reference_run(run)
                
                # 检查是否为上标数字
                if run.font.superscript:
                    self._format_superscript_footnote(run)
    
    def _contains_footnote_marker(self, text):
        """检查文本是否包含脚注标记"""
        return bool(self.footnote_pattern.search(text))
    
    def _format_reference_run(self, run):
        """格式化包含脚注标记的run"""
        # 查找脚注标记
        text = run.text
        match = self.footnote_pattern.search(text)
        
        if match:
            # 获取标记前后的文本
            marker = match.group()
            before = text[:match.start()]
            after = text[match.end():]
            
            # 重新构建run
            run.text = before
            
            # 添加格式化的脚注标记
            if before:
                # 如果前面有文本，需要新建run
                para = run._element.getparent()
                new_run = run._element.addnext(OxmlElement('w:r'))
                new_run_obj = run.__class__(new_run, run._parent)
                new_run_obj.text = marker
                new_run_obj.font.superscript = True
                new_run_obj.font.size = Pt(9)  # 小五号
                new_run_obj.font.name = '宋体'
                new_run_obj._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                # 直接修改当前run
                run.text = marker
                run.font.superscript = True
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 如果后面还有文本，添加新run
            if after:
                para = run._element.getparent()
                after_run = run._element.addnext(OxmlElement('w:r'))
                after_run_obj = run.__class__(after_run, run._parent)
                after_run_obj.text = after
    
    def _format_superscript_footnote(self, run):
        """格式化上标脚注"""
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)  # 小五号
    
    def _format_footnote_content(self, doc):
        """
        格式化脚注内容区域
        注：python-docx对脚注内容的直接访问有限，这里提供基本实现
        """
        # 查找可能的脚注内容（通常在页面底部）
        for para in doc.paragraphs:
            if self._is_footnote_content(para):
                self._format_footnote_paragraph(para)
    
    def _is_footnote_content(self, paragraph):
        """判断段落是否为脚注内容"""
        text = paragraph.text.strip()
        
        # 检查是否以脚注标记开头
        if self.footnote_pattern.match(text):
            # 检查字体大小（脚注通常较小）
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.font.size and first_run.font.size <= Pt(10):
                    return True
        
        # 检查是否在页面底部（通过分隔线等特征）
        # 注：这需要更复杂的文档结构分析
        
        return False
    
    def _format_footnote_paragraph(self, paragraph):
        """格式化脚注段落"""
        # 设置段落格式
        paragraph.paragraph_format.line_spacing = Pt(12)  # 行距12pt
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        
        # 格式化所有run
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)  # 小五号
    
    def add_footnote_separator(self, doc):
        """添加脚注分隔线"""
        # 在需要的位置添加分隔线
        # 注：这通常需要在节的设置中完成
        pass
    
    def convert_endnotes_to_footnotes(self, doc):
        """将尾注转换为脚注格式"""
        # 查找可能的尾注区域
        endnotes_start = self._find_endnotes_section(doc)
        
        if endnotes_start >= 0:
            # 收集尾注内容
            endnotes = self._collect_endnotes(doc, endnotes_start)
            
            # 将尾注转换为脚注格式
            self._convert_to_footnotes(doc, endnotes)
    
    def _find_endnotes_section(self, doc):
        """查找尾注部分的起始位置"""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if any(keyword in text for keyword in ['注释', '尾注', 'notes', 'endnotes']):
                # 检查是否为标题格式
                if para.runs and para.runs[0].font.bold:
                    return i
        return -1
    
    def _collect_endnotes(self, doc, start_idx):
        """收集尾注内容"""
        endnotes = []
        
        for i in range(start_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # 检查是否为尾注项
            if self.footnote_pattern.match(text):
                endnotes.append({
                    'number': self._extract_note_number(text),
                    'content': text,
                    'paragraph': para
                })
            elif text and i > start_idx + 1:
                # 可能是上一个尾注的续行
                if endnotes:
                    endnotes[-1]['content'] += ' ' + text
        
        return endnotes
    
    def _extract_note_number(self, text):
        """提取注释编号"""
        match = self.footnote_pattern.match(text)
        if match:
            number_text = match.group()
            # 提取数字
            number_match = re.search(r'\d+', number_text)
            if number_match:
                return int(number_match.group())
        return 0
    
    def _convert_to_footnotes(self, doc, endnotes):
        """将收集的尾注转换为脚注格式"""
        # 这里提供概念性实现
        # 实际转换需要更复杂的Word文档操作
        for note in endnotes:
            # 格式化尾注段落为脚注样式
            para = note['paragraph']
            self._format_footnote_paragraph(para)
    
    def ensure_page_restart_numbering(self, doc):
        """
        确保脚注在每页重新编号
        注：这通常需要在Word的节设置中完成
        """
        # 为每个节设置脚注重新编号
        for section in doc.sections:
            # 这需要访问底层XML来设置
            # Word中的脚注设置存储在sectPr中
            sectPr = section._sectPr
            
            # 查找或创建footnotePr元素
            footnotePr = sectPr.find(qn('w:footnotePr'))
            if footnotePr is None:
                footnotePr = OxmlElement('w:footnotePr')
                sectPr.append(footnotePr)
            
            # 设置numRestart属性为eachPage
            numRestart = footnotePr.find(qn('w:numRestart'))
            if numRestart is None:
                numRestart = OxmlElement('w:numRestart')
                footnotePr.append(numRestart)
            
            numRestart.set(qn('w:val'), 'eachPage')


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加一些带脚注的内容
    p1 = doc.add_paragraph("这是一段包含脚注的文本")
    run1 = p1.add_run("[1]")
    run1.font.superscript = True
    p1.add_run("，用于测试脚注格式化功能。")
    
    p2 = doc.add_paragraph("另一段文本包含多个脚注")
    run2 = p2.add_run("①")
    p2.add_run("和")
    run3 = p2.add_run("②")
    p2.add_run("。")
    
    # 添加脚注内容（模拟）
    doc.add_paragraph()  # 空行
    doc.add_paragraph("―" * 50)  # 分隔线
    
    fn1 = doc.add_paragraph("[1] 这是第一个脚注的内容。")
    fn1.runs[0].font.size = Pt(9)
    
    fn2 = doc.add_paragraph("① 这是第二个脚注的内容。")
    fn2.runs[0].font.size = Pt(9)
    
    fn3 = doc.add_paragraph("② 这是第三个脚注的内容。")
    fn3.runs[0].font.size = Pt(9)
    
    # 测试格式化
    formatter = FootnoteFormatter()
    formatter.format_footnotes(doc)
    
    # 保存测试文档
    doc.save('test_footnotes.docx')
    print("测试文档已生成：test_footnotes.docx")
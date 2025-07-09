#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
主格式化器模块
集成所有格式化功能，实现完整的论文格式化
"""

from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
import sys
from datetime import datetime
import importlib.util

# 导入所有格式化模块
try:
    from .cover_generator import CoverGenerator, CommitmentGenerator
    from .page_number_handler import PageNumberHandler
    from .keyword_formatter import KeywordFormatter
    from .figure_table_handler import FigureTableHandler
    from .footnote_formatter import FootnoteFormatter
    from .math_formatter import MathFormatter
    from .toc_generator import TOCGenerator
    from .acknowledgment_formatter import AcknowledgmentFormatter
    from .appendix_handler import AppendixHandler
    from .document_reorganizer import DocumentReorganizer
    from .font_manager import FontManager
    from .spacing_manager import SpacingManager
    from .header_handler import HeaderHandler
    from .basic_formatters import (
        HeadingFormatter, ParagraphFormatter,
        AbstractFormatter, ReferenceFormatter
    )
except ImportError:
    # 如果相对导入失败，尝试绝对导入
    from cover_generator import CoverGenerator, CommitmentGenerator
    from page_number_handler import PageNumberHandler
    from keyword_formatter import KeywordFormatter
    from figure_table_handler import FigureTableHandler
    from footnote_formatter import FootnoteFormatter
    from math_formatter import MathFormatter
    from toc_generator import TOCGenerator
    from acknowledgment_formatter import AcknowledgmentFormatter
    from appendix_handler import AppendixHandler
    from document_reorganizer import DocumentReorganizer
    from font_manager import FontManager
    from spacing_manager import SpacingManager
    from header_handler import HeaderHandler
    from basic_formatters import (
        HeadingFormatter, ParagraphFormatter,
        AbstractFormatter, ReferenceFormatter
    )

# 创建简单的文档分析器
class DocumentAnalyzer:
    """文档结构分析器"""
    def __init__(self, document):
        self.document = document
    
    def analyze(self):
        """分析文档结构"""
        structure = {
            'abstract_cn': -1,
            'abstract_en': -1,
            'toc': -1,
            'main_start': -1,
            'references': -1
        }
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            
            if text == '摘要' and structure['abstract_cn'] == -1:
                structure['abstract_cn'] = i
            elif text == 'Abstract' and structure['abstract_en'] == -1:
                structure['abstract_en'] = i
            elif text == '目录' and structure['toc'] == -1:
                structure['toc'] = i
            elif '第一章' in text or '第1章' in text:
                if structure['main_start'] == -1:
                    structure['main_start'] = i
            elif text == '参考文献' and structure['references'] == -1:
                structure['references'] = i
        
        return structure


class CompleteThesisFormatter:
    """完整的论文格式化器"""
    
    def __init__(self, document_path=None, config=None):
        """
        初始化格式化器
        
        Args:
            document_path: Word文档路径（可选，支持后续设置）
            config: 配置字典
        """
        self.document_path = document_path
        self.document = None
        self.config = config or self._get_default_config()
        self.format_options = {}  # 支持GUI选项
        
        # 如果提供了文档路径，立即加载
        if document_path:
            self.load_document(document_path)
    
    def load_document(self, document_path):
        """加载文档"""
        self.document_path = document_path
        self.document = Document(document_path)
        
        # 初始化所有处理器
        self._init_processors()
        
        # 线程池
        self.executor = ThreadPoolExecutor(max_workers=12)
        
    def _get_default_config(self):
        """获取默认配置"""
        return {
            'generate_cover': True,
            'generate_commitment': True,
            'format_keywords': True,
            'format_figures_tables': True,
            'format_footnotes': True,
            'format_math': True,
            'update_toc': True,
            'format_acknowledgment': True,
            'format_appendix': True,
            'setup_page_numbers': True,
            'reorder_document': True,
            'basic_formatting': True
        }
    
    def _init_processors(self):
        """初始化所有处理器"""
        # 文档结构分析器
        self.analyzer = DocumentAnalyzer(self.document)
        
        # 基础格式化器
        self.heading_formatter = HeadingFormatter(self.document, {})
        self.paragraph_formatter = ParagraphFormatter(self.document, {})
        self.abstract_formatter = AbstractFormatter(self.document, {})
        self.reference_formatter = ReferenceFormatter(self.document, {})
        
        # 新增格式化器
        self.cover_generator = CoverGenerator()
        self.commitment_generator = CommitmentGenerator()
        self.page_handler = PageNumberHandler(self.document)
        self.keyword_formatter = KeywordFormatter()
        self.figure_table_handler = FigureTableHandler()
        self.footnote_formatter = FootnoteFormatter()
        self.math_formatter = MathFormatter()
        self.toc_generator = TOCGenerator()
        self.acknowledgment_formatter = AcknowledgmentFormatter()
        self.appendix_handler = AppendixHandler()
        self.document_reorganizer = DocumentReorganizer()
        
        # 新增的核心格式化器
        self.font_manager = FontManager()
        self.spacing_manager = SpacingManager()
        self.header_handler = HeaderHandler()
    
    def format_document(self, input_file=None, output_file=None, thesis_info=None, progress_callback=None):
        """
        执行完整的文档格式化
        
        Args:
            input_file: 输入文件路径（可选，如果未初始化时提供）
            output_file: 输出文件路径（可选）
            thesis_info: 论文信息字典
            progress_callback: 进度回调函数
        
        Returns:
            bool: 格式化是否成功
        """
        # 如果提供了输入文件，加载它
        if input_file:
            self.load_document(input_file)
        
        # 如果设置了format_options，使用它覆盖config
        if self.format_options:
            for key, value in self.format_options.items():
                if key in ['cover', 'commitment', 'page_number', 'keywords', 
                          'figures_tables', 'footnotes', 'math', 'toc', 
                          'acknowledgment', 'appendix', 'reorganize', 'basic']:
                    # 映射GUI选项到内部配置
                    config_key = self._map_option_to_config(key)
                    self.config[config_key] = value
        
        # 保存原始的format_document方法逻辑
        return self._format_document_impl(thesis_info, progress_callback, output_file)
    
    def _map_option_to_config(self, option_key):
        """映射GUI选项到配置键"""
        mapping = {
            'cover': 'generate_cover',
            'commitment': 'generate_commitment',
            'page_number': 'setup_page_numbers',
            'keywords': 'format_keywords',
            'figures_tables': 'format_figures_tables',
            'footnotes': 'format_footnotes',
            'math': 'format_math',
            'toc': 'update_toc',
            'acknowledgment': 'format_acknowledgment',
            'appendix': 'format_appendix',
            'reorganize': 'reorder_document',
            'basic': 'basic_formatting'
        }
        return mapping.get(option_key, option_key)
    
    def _format_document_impl(self, thesis_info=None, progress_callback=None, output_file=None):
        """
        执行完整的文档格式化
        
        Args:
            thesis_info: 论文信息字典
            progress_callback: 进度回调函数
        
        Returns:
            bool: 格式化是否成功
        """
        try:
            # 默认论文信息
            if thesis_info is None:
                thesis_info = {
                    'title': '基于深度学习的图像识别研究',
                    'major': '计算机科学与技术',
                    'class': '计科1901',
                    'student_id': '20190001',
                    'name': '张三',
                    'advisor': '李教授',
                    'date': datetime.now().strftime('%Y年%m月')
                }
            
            # 分析文档结构
            if progress_callback:
                progress_callback("分析文档结构...", 5)
            structure = self.analyzer.analyze()
            
            # 创建并行任务列表
            tasks = self._create_tasks(structure, thesis_info)
            
            # 执行并行任务
            completed = 0
            total = len(tasks)
            
            for future in as_completed(tasks):
                completed += 1
                if progress_callback:
                    progress = int((completed / total) * 90) + 5
                    progress_callback(f"正在格式化... ({completed}/{total})", progress)
                
                # 检查任务结果
                try:
                    result = future.result()
                except Exception as e:
                    print(f"任务执行失败: {e}")
            
            # 执行需要顺序处理的任务
            if progress_callback:
                progress_callback("执行后处理...", 95)
            self._post_process(structure, thesis_info)
            
            # 保存文档
            if progress_callback:
                progress_callback("保存文档...", 98)
            
            # 使用指定的输出文件名或自动生成
            if output_file:
                self.document.save(output_file)
                print(f"文档已保存: {output_file}")
            else:
                self._save_document()
            
            if progress_callback:
                progress_callback("格式化完成！", 100)
            
            return True
            
        except Exception as e:
            print(f"格式化失败: {e}")
            return False
    
    def _create_tasks(self, structure, thesis_info):
        """创建并行任务列表"""
        tasks = []
        
        # 基础格式化任务
        if self.config.get('basic_formatting', True):
            tasks.append(
                self.executor.submit(self.heading_formatter.format_headings, structure)
            )
            tasks.append(
                self.executor.submit(self.paragraph_formatter.format_paragraphs, structure)
            )
            tasks.append(
                self.executor.submit(self.abstract_formatter.format_abstract, structure)
            )
            tasks.append(
                self.executor.submit(self.reference_formatter.format_references, structure)
            )
        
        # 新增格式化任务
        if self.config.get('format_keywords', True):
            tasks.append(
                self.executor.submit(self.keyword_formatter.format_keywords, 
                                   self.document, structure)
            )
        
        if self.config.get('format_figures_tables', True):
            tasks.append(
                self.executor.submit(self.figure_table_handler.process_figures_and_tables,
                                   self.document, structure)
            )
        
        if self.config.get('format_footnotes', True):
            tasks.append(
                self.executor.submit(self.footnote_formatter.format_footnotes,
                                   self.document)
            )
        
        if self.config.get('format_math', True):
            tasks.append(
                self.executor.submit(self.math_formatter.format_math_formulas,
                                   self.document, structure)
            )
        
        if self.config.get('format_acknowledgment', True):
            tasks.append(
                self.executor.submit(self.acknowledgment_formatter.format_acknowledgment,
                                   self.document, structure)
            )
        
        if self.config.get('format_appendix', True):
            tasks.append(
                self.executor.submit(self.appendix_handler.process_appendix,
                                   self.document, structure)
            )
        
        return tasks
    
    def _post_process(self, structure, thesis_info):
        """执行需要顺序处理的后处理任务"""
        # 1. 生成封面和承诺书（需要在最前面）
        if self.config.get('generate_cover', True):
            self.cover_generator.generate_cover(self.document, thesis_info)
        
        if self.config.get('generate_commitment', True):
            self.commitment_generator.generate_commitment(self.document)
        
        # 2. 重新分析文档结构（因为添加了新内容）
        structure = self.analyzer.analyze()
        
        # 3. 应用字体管理器 - 修正字体字号问题
        if self.config.get('basic_formatting', True):
            self._apply_font_formatting(structure)
        
        # 4. 应用行间距管理器 - 分区域设置行间距
        if self.config.get('basic_formatting', True):
            self.spacing_manager.process_document_spacing(self.document, structure)
        
        # 5. 添加页眉（仅正文部分）
        if self.config.get('setup_page_numbers', True):
            self.header_handler.add_page_header_to_document(self.document, structure)
        
        # 6. 设置页码系统
        if self.config.get('setup_page_numbers', True):
            self.page_handler.setup_page_numbers(structure)
        
        # 7. 更新目录（需要在所有内容确定后）
        if self.config.get('update_toc', True):
            self.toc_generator.update_toc(self.document, structure)
        
        # 8. 文档结构重组（如果需要）
        if self.config.get('reorder_document', True):
            self._reorder_document(structure)
    
    def _reorder_document(self, structure):
        """
        按照标准顺序重新组织文档
        标准顺序：封面 → 承诺书 → 中文摘要 → 英文摘要 → 目录 → 正文 → 参考文献 → 致谢 → 附录
        """
        # 使用DocumentReorganizer重组文档
        new_doc = self.document_reorganizer.reorganize_document(self.document, structure)
        
        # 更新当前文档
        self.document = new_doc
    
    def _save_document(self):
        """保存格式化后的文档"""
        # 生成新文件名
        base_name = os.path.splitext(self.document_path)[0]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"{base_name}_formatted_{timestamp}.docx"
        
        # 保存文档
        self.document.save(output_path)
        print(f"文档已保存: {output_path}")
        
        return output_path
    
    def get_format_report(self):
        """生成格式化报告"""
        report = []
        report.append("="*50)
        report.append("论文格式化报告")
        report.append("="*50)
        
        # 检查各项格式
        checks = [
            ("封面", self._check_cover()),
            ("承诺书", self._check_commitment()),
            ("页码系统", self._check_page_numbers()),
            ("关键词格式", self._check_keywords()),
            ("图表编号", self._check_figures_tables()),
            ("目录", self._check_toc()),
            ("致谢", self._check_acknowledgment())
        ]
        
        for item, status in checks:
            status_text = "✓ 已完成" if status else "✗ 未完成"
            report.append(f"{item}: {status_text}")
        
        report.append("="*50)
        
        return "\n".join(report)
    
    def _check_cover(self):
        """检查是否有封面"""
        # 简单检查第一页是否包含学校名称
        if self.document.paragraphs:
            first_para = self.document.paragraphs[0].text
            return "江西财经大学" in first_para
        return False
    
    def _check_commitment(self):
        """检查是否有承诺书"""
        for para in self.document.paragraphs[:20]:  # 检查前20个段落
            if "诚信承诺书" in para.text or "承诺书" in para.text:
                return True
        return False
    
    def _check_page_numbers(self):
        """检查页码设置"""
        # 检查是否有多个节
        return len(self.document.sections) > 1
    
    def _check_keywords(self):
        """检查关键词格式"""
        for para in self.document.paragraphs:
            if "[关键词]" in para.text or "[Keywords]" in para.text:
                return True
        return False
    
    def _check_figures_tables(self):
        """检查图表编号"""
        for para in self.document.paragraphs:
            if "图1.1" in para.text or "表1.1" in para.text:
                return True
        return False
    
    def _check_toc(self):
        """检查目录"""
        for para in self.document.paragraphs:
            if para.text.strip() == "目录":
                # 检查后续是否有目录内容
                return True
        return False
    
    def _check_acknowledgment(self):
        """检查致谢"""
        for para in self.document.paragraphs:
            if "致谢" in para.text and len(para.text) < 10:
                return True
        return False
    
    def _apply_font_formatting(self, structure):
        """应用字体格式化"""
        try:
            for paragraph in self.document.paragraphs:
                text = paragraph.text.strip()
                
                # 处理摘要标题
                if text == '摘要':
                    for run in paragraph.runs:
                        self.font_manager.apply_font_style(run, 'abstract_title_cn', 'cn')
                    self.font_manager.apply_paragraph_style(paragraph, 'abstract_title_cn')
                
                elif text == 'Abstract':
                    for run in paragraph.runs:
                        self.font_manager.apply_font_style(run, 'abstract_title_en', 'en')
                    self.font_manager.apply_paragraph_style(paragraph, 'abstract_title_en')
                
                # 处理标题
                elif self._is_heading_paragraph(paragraph):
                    level = self._get_heading_level(text)
                    style_name = f'heading_{level}' if level <= 3 else 'heading_3'
                    self.font_manager.format_mixed_text(paragraph, style_name)
                
                # 处理正文
                elif text and not self._is_special_paragraph(text):
                    self.font_manager.format_mixed_text(paragraph, 'main_text')
                
        except Exception as e:
            self.logger.error(f"字体格式化失败: {e}")
    
    def _is_heading_paragraph(self, paragraph):
        """判断是否为标题段落"""
        if not paragraph.runs:
            return False
        
        # 检查是否有加粗
        if paragraph.runs[0].font.bold:
            text = paragraph.text.strip()
            # 检查标题模式
            import re
            if re.match(r'^第[一二三四五六七八九十\d]+[章节]', text):
                return True
            if re.match(r'^\d+\.\d*\s+', text):
                return True
        return False
    
    def _get_heading_level(self, text):
        """获取标题级别"""
        import re
        if re.match(r'^第[一二三四五六七八九十\d]+[章节]', text):
            return 1
        elif re.match(r'^\d+\.\d+\s+', text):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+\s+', text):
            return 3
        return 1
    
    def _is_special_paragraph(self, text):
        """判断是否为特殊段落（不需要正文格式化）"""
        special_keywords = ['摘要', 'Abstract', '目录', '参考文献', '致谢', '附录']
        return any(keyword in text for keyword in special_keywords)


def main():
    """测试主函数"""
    import tkinter as tk
    from tkinter import filedialog, messagebox
    from tkinter import ttk
    
    def select_file():
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            file_path.set(filename)
    
    def format_document():
        if not file_path.get():
            messagebox.showerror("错误", "请先选择文档")
            return
        
        # 收集论文信息
        thesis_info = {
            'title': title_var.get() or '毕业论文题目',
            'major': major_var.get() or '专业名称',
            'class': class_var.get() or '班级',
            'student_id': student_id_var.get() or '学号',
            'name': name_var.get() or '姓名',
            'advisor': advisor_var.get() or '指导教师',
            'date': date_var.get() or datetime.now().strftime('%Y年%m月')
        }
        
        # 更新进度
        def update_progress(message, value):
            progress_var.set(value)
            status_label.config(text=message)
            root.update()
        
        # 执行格式化
        formatter = CompleteThesisFormatter(file_path.get())
        success = formatter.format_document(thesis_info, update_progress)
        
        if success:
            # 显示报告
            report = formatter.get_format_report()
            messagebox.showinfo("完成", f"格式化完成！\n\n{report}")
        else:
            messagebox.showerror("错误", "格式化失败")
    
    # 创建GUI
    root = tk.Tk()
    root.title("毕业论文一键格式化工具 - 完整版")
    root.geometry("600x700")
    
    # 文件选择
    file_frame = ttk.LabelFrame(root, text="选择文档", padding=10)
    file_frame.pack(fill='x', padx=20, pady=10)
    
    file_path = tk.StringVar()
    ttk.Entry(file_frame, textvariable=file_path, width=50).pack(side='left', padx=5)
    ttk.Button(file_frame, text="浏览", command=select_file).pack(side='left')
    
    # 论文信息
    info_frame = ttk.LabelFrame(root, text="论文信息", padding=10)
    info_frame.pack(fill='x', padx=20, pady=10)
    
    # 创建输入字段
    fields = [
        ('论文题目:', 'title_var'),
        ('专业:', 'major_var'),
        ('班级:', 'class_var'),
        ('学号:', 'student_id_var'),
        ('姓名:', 'name_var'),
        ('指导教师:', 'advisor_var'),
        ('日期:', 'date_var')
    ]
    
    # 创建变量
    title_var = tk.StringVar()
    major_var = tk.StringVar()
    class_var = tk.StringVar()
    student_id_var = tk.StringVar()
    name_var = tk.StringVar()
    advisor_var = tk.StringVar()
    date_var = tk.StringVar(value=datetime.now().strftime('%Y年%m月'))
    
    vars_dict = {
        'title_var': title_var,
        'major_var': major_var,
        'class_var': class_var,
        'student_id_var': student_id_var,
        'name_var': name_var,
        'advisor_var': advisor_var,
        'date_var': date_var
    }
    
    for label, var_name in fields:
        frame = ttk.Frame(info_frame)
        frame.pack(fill='x', pady=2)
        ttk.Label(frame, text=label, width=10).pack(side='left')
        ttk.Entry(frame, textvariable=vars_dict[var_name], width=40).pack(side='left', padx=5)
    
    # 进度条
    progress_frame = ttk.Frame(root)
    progress_frame.pack(fill='x', padx=20, pady=10)
    
    progress_var = tk.IntVar()
    progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100)
    progress_bar.pack(fill='x')
    
    status_label = ttk.Label(progress_frame, text="准备就绪")
    status_label.pack()
    
    # 格式化按钮
    ttk.Button(root, text="开始格式化", command=format_document, 
               style='Accent.TButton').pack(pady=20)
    
    root.mainloop()


if __name__ == "__main__":
    main()
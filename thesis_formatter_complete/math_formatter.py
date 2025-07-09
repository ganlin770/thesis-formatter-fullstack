#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
数学公式格式化模块
实现公式居中显示，编号右对齐
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
import re

class MathFormatter:
    """数学公式格式化器"""
    
    def __init__(self):
        self.formula_counter = 0
        self.chapter_counters = {}  # 章节 -> 公式计数
        self.current_chapter = 1
        
        # 公式识别模式
        self.formula_patterns = [
            r'^\s*[a-zA-Z]\s*=',  # 以字母=开头
            r'=.*[+\-*/]',        # 包含数学运算符
            r'\\[a-zA-Z]+',       # LaTeX命令
            r'∫|∑|∏|√|∞',        # 数学符号
            r'\d+\s*[+\-*/]\s*\d+',  # 简单算式
        ]
    
    def format_math_formulas(self, doc, structure):
        """
        格式化文档中的数学公式
        
        Args:
            doc: Document对象
            structure: 文档结构字典
        """
        main_start = structure.get('main_start', 0)
        
        for i in range(main_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            
            # 检测章节
            chapter_num = self._detect_chapter(para)
            if chapter_num:
                self.current_chapter = chapter_num
                if chapter_num not in self.chapter_counters:
                    self.chapter_counters[chapter_num] = 0
            
            # 检测并格式化公式
            if self._is_formula(para):
                self._format_formula(para)
    
    def _detect_chapter(self, paragraph):
        """检测章节编号"""
        text = paragraph.text.strip()
        
        # 章节标题模式
        patterns = [
            r'^第([一二三四五六七八九十\d]+)[章节]',
            r'^(\\d+)[.\\s]+\\S',
            r'^Chapter\\s+(\\d+)'
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                chapter_str = match.group(1)
                # 转换中文数字
                if chapter_str in '一二三四五六七八九十':
                    return self._chinese_to_arabic(chapter_str)
                else:
                    return int(chapter_str)
        
        return None
    
    def _chinese_to_arabic(self, chinese_num):
        """中文数字转阿拉伯数字"""
        num_map = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10
        }
        return num_map.get(chinese_num, 1)
    
    def _is_formula(self, paragraph):
        """判断段落是否为数学公式"""
        text = paragraph.text.strip()
        
        # 空段落不是公式
        if not text:
            return False
        
        # 太长的段落不太可能是独立公式
        if len(text) > 200:
            return False
        
        # 检查是否匹配公式模式
        for pattern in self.formula_patterns:
            if re.search(pattern, text):
                # 额外检查：是否为独立段落（前后可能有空行）
                if self._is_standalone_paragraph(paragraph):
                    return True
        
        # 检查是否包含常见数学符号
        math_symbols = ['=', '+', '-', '*', '/', '^', '(', ')', '[', ']', 
                       'α', 'β', 'γ', 'δ', 'θ', 'λ', 'μ', 'π', 'σ', 'φ']
        
        symbol_count = sum(1 for symbol in math_symbols if symbol in text)
        
        # 如果包含多个数学符号，可能是公式
        if symbol_count >= 2 and len(text) < 100:
            return True
        
        return False
    
    def _is_standalone_paragraph(self, paragraph):
        """判断是否为独立段落（可能是公式）"""
        # 检查对齐方式
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        
        # 检查段落前后间距
        if paragraph.paragraph_format.space_before and paragraph.paragraph_format.space_before > Pt(6):
            return True
        
        return False
    
    def _format_formula(self, paragraph):
        """格式化数学公式段落"""
        # 增加公式计数
        if self.current_chapter not in self.chapter_counters:
            self.chapter_counters[self.current_chapter] = 0
        self.chapter_counters[self.current_chapter] += 1
        
        # 生成公式编号
        formula_number = f"({self.current_chapter}.{self.chapter_counters[self.current_chapter]})"
        
        # 获取原始公式文本
        formula_text = paragraph.text.strip()
        
        # 检查是否已有编号
        existing_number_pattern = r'\s*\(\d+\.\d+\)\s*$'
        if re.search(existing_number_pattern, formula_text):
            # 移除现有编号
            formula_text = re.sub(existing_number_pattern, '', formula_text).strip()
        
        # 清空段落
        paragraph.clear()
        
        # 设置段落格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = Pt(22)
        
        # 添加制表位（用于右对齐编号）
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.clear_all()
        # 在页面右边距附近添加右对齐制表位
        tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT)
        
        # 添加公式文本
        run_formula = paragraph.add_run(formula_text)
        run_formula.font.name = 'Times New Roman'
        run_formula.font.size = Pt(12)
        
        # 添加制表符
        paragraph.add_run('\t')
        
        # 添加公式编号
        run_number = paragraph.add_run(formula_number)
        run_number.font.name = 'Times New Roman'
        run_number.font.size = Pt(12)
    
    def format_inline_formulas(self, doc):
        """格式化行内公式"""
        for para in doc.paragraphs:
            if not self._is_formula(para):  # 不是独立公式
                self._format_inline_math(para)
    
    def _format_inline_math(self, paragraph):
        """格式化段落中的行内数学符号"""
        # 查找需要斜体的变量
        for run in paragraph.runs:
            text = run.text
            
            # 单字母变量（在数学上下文中）
            if self._contains_math_variables(text):
                self._format_math_run(run)
    
    def _contains_math_variables(self, text):
        """检查是否包含数学变量"""
        # 匹配独立的单字母变量
        variable_pattern = r'\b[a-zA-Z]\b'
        
        # 检查上下文
        math_context = ['=', '+', '-', '*', '/', '^', '(', ')']
        
        if any(symbol in text for symbol in math_context):
            if re.search(variable_pattern, text):
                return True
        
        return False
    
    def _format_math_run(self, run):
        """格式化包含数学内容的run"""
        # 对于英文字母，使用Times New Roman斜体
        if re.search(r'[a-zA-Z]', run.text):
            run.font.name = 'Times New Roman'
            run.font.italic = True
    
    def add_formula_if_needed(self, paragraph, formula_text):
        """
        添加格式化的公式到段落
        
        Args:
            paragraph: 目标段落
            formula_text: 公式文本
        """
        # 清空段落
        paragraph.clear()
        
        # 设置为公式格式
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        
        # 添加公式
        run = paragraph.add_run(formula_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # 标记为公式段落，后续可以添加编号
        # 这里可以添加自定义属性或样式
    
    def update_formula_references(self, doc):
        """更新文档中的公式引用"""
        # 存储公式编号映射
        formula_map = {}  # "公式1" -> "公式(1.1)"
        
        # 第一遍：收集所有公式编号
        for para in doc.paragraphs:
            text = para.text
            
            # 查找公式编号
            number_match = re.search(r'\((\d+)\.(\d+)\)', text)
            if number_match and self._is_formula(para):
                chapter = number_match.group(1)
                number = number_match.group(2)
                
                # 创建映射
                old_ref = f"公式{number}"
                new_ref = f"公式({chapter}.{number})"
                formula_map[old_ref] = new_ref
                
                # 也映射简单格式
                formula_map[f"式{number}"] = new_ref
                formula_map[f"式({number})"] = new_ref
        
        # 第二遍：更新所有引用
        for para in doc.paragraphs:
            if self._is_formula(para):
                continue  # 跳过公式本身
            
            text = para.text
            changed = False
            
            # 替换引用
            for old_ref, new_ref in formula_map.items():
                if old_ref in text:
                    text = text.replace(old_ref, new_ref)
                    changed = True
            
            # 如果有改变，更新段落
            if changed:
                # 保存原有格式
                if para.runs:
                    # 简单情况：保持第一个run的格式
                    font_name = para.runs[0].font.name
                    font_size = para.runs[0].font.size
                    
                    para.clear()
                    run = para.add_run(text)
                    if font_name:
                        run.font.name = font_name
                        if font_name in ['宋体', '楷体', '黑体', '仿宋']:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if font_size:
                        run.font.size = font_size


if __name__ == "__main__":
    # 测试代码
    doc = Document()
    
    # 添加章节
    p1 = doc.add_paragraph("第一章 数学模型")
    p1.runs[0].font.bold = True
    
    doc.add_paragraph("下面是一些数学公式的例子：")
    
    # 添加公式
    formula1 = doc.add_paragraph("E = mc²")
    formula1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    formula2 = doc.add_paragraph("a² + b² = c²")
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("根据公式1可知，能量与质量的关系...")
    
    # 添加第二章
    p2 = doc.add_paragraph("第二章 算法分析")
    p2.runs[0].font.bold = True
    
    formula3 = doc.add_paragraph("f(x) = ∫ₐᵇ g(t)dt")
    formula3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("时间复杂度为 O(n²)，其中 n 是输入规模。")
    
    # 测试格式化
    formatter = MathFormatter()
    structure = {'main_start': 0}
    
    formatter.format_math_formulas(doc, structure)
    formatter.format_inline_formulas(doc)
    formatter.update_formula_references(doc)
    
    # 保存测试文档
    doc.save('test_math_formulas.docx')
    print("测试文档已生成：test_math_formulas.docx")